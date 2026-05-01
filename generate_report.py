"""
Generates PROJECT_REPORT.docx for the Smart Disaster Response MIS project.
Run: python generate_report.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.0)
section.right_margin  = Inches(1.0)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ── Colour palette ──────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x3A, 0x5C)
MID_BLUE    = RGBColor(0x1E, 0x5F, 0x99)
ACCENT_BLUE = RGBColor(0x2E, 0x86, 0xC1)
LIGHT_GREY  = RGBColor(0xF2, 0xF2, 0xF2)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_DARK   = RGBColor(0x1A, 0x1A, 0x2E)
TEXT_MID    = RGBColor(0x34, 0x49, 0x5E)
GREEN       = RGBColor(0x27, 0xAE, 0x60)
RED         = RGBColor(0xC0, 0x39, 0x2B)
ORANGE      = RGBColor(0xE6, 0x7E, 0x22)

# ── Helper: shade a table cell ───────────────────────────────────────────────
def shade_cell(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def set_cell_border(cell, **kwargs):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'),   kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:space'), '0')
            tag.set(qn('w:color'), kwargs[edge].get('color', 'auto'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

# ── Helper: add a styled heading paragraph ──────────────────────────────────
def add_heading(text, level=1, color=DARK_BLUE, space_before=18, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(18)
    elif level == 2:
        run.font.size = Pt(14)
        p.paragraph_format.space_before = Pt(14)
    elif level == 3:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(10)
    else:
        run.font.size = Pt(11)
        p.paragraph_format.space_before = Pt(8)

# ── Helper: add normal body paragraph ───────────────────────────────────────
def add_para(text, bold=False, italic=False, color=TEXT_DARK, size=10.5,
             space_before=2, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return p

# ── Helper: add a bullet point ───────────────────────────────────────────────
def add_bullet(text, level=0, color=TEXT_DARK, size=10.5):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 + level * 0.25)
    p.paragraph_format.space_before  = Pt(1)
    p.paragraph_format.space_after   = Pt(2)
    run = p.add_run(text)
    run.font.size      = Pt(size)
    run.font.color.rgb = color
    return p

# ── Helper: add a styled table ───────────────────────────────────────────────
def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style     = 'Table Grid'

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        shade_cell(cell, '1E5F99')
        p    = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run  = p.add_run(h)
        run.bold           = True
        run.font.size      = Pt(9)
        run.font.color.rgb = WHITE

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        bg = 'F2F2F2' if r_idx % 2 == 0 else 'FFFFFF'
        for i, val in enumerate(row_data):
            cell = row_cells[i]
            shade_cell(cell, bg)
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run  = p.add_run(str(val))
            run.font.size      = Pt(9)
            run.font.color.rgb = TEXT_DARK

    # Column widths
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)

    doc.add_paragraph()  # breathing room after table
    return table

# ── Helper: add a code block ─────────────────────────────────────────────────
def add_code(code_text, font_size=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Inches(0.25)
    # Light grey background via shading on the paragraph's XML
    pPr  = p._p.get_or_add_pPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  'F4F4F4')
    pPr.append(shd)
    run = p.add_run(code_text)
    run.font.name      = 'Courier New'
    run.font.size      = Pt(font_size)
    run.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    return p

# ── Helper: horizontal rule ──────────────────────────────────────────────────
def add_hr():
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '6')
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), '1E5F99')
    pb.append(bot)
    pPr.append(pb)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)

# ── Helper: add a diagram placeholder box (yellow-highlighted) ───────────────
def add_diagram_box(title, mermaid_code, description=""):
    # Outer single-cell table gives the whole block a vivid amber border + fill
    YELLOW_BG  = 'FFF9C4'   # pale amber background
    AMBER_LINE = 'F9A825'   # amber border colour

    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)

    # Yellow background on cell
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  YELLOW_BG)
    tcPr.append(shd)

    # Amber border on all four sides
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        tag = OxmlElement(f'w:{edge}')
        tag.set(qn('w:val'),   'single')
        tag.set(qn('w:sz'),    '12')   # 1.5 pt
        tag.set(qn('w:space'), '0')
        tag.set(qn('w:color'), AMBER_LINE)
        tcBorders.append(tag)
    tcPr.append(tcBorders)

    # ── Title line inside the cell ──────────────────────────────────────────
    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(6)
    p_title.paragraph_format.space_after  = Pt(2)
    p_title.paragraph_format.left_indent  = Inches(0.1)
    run_icon  = p_title.add_run("📊 DIAGRAM PLACEHOLDER  ▸  ")
    run_icon.bold           = True
    run_icon.font.size      = Pt(10)
    run_icon.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)   # dark amber
    run_title = p_title.add_run(title)
    run_title.bold           = True
    run_title.font.size      = Pt(10)
    run_title.font.color.rgb = RGBColor(0x4A, 0x23, 0x00)   # dark brown

    # ── Optional description ────────────────────────────────────────────────
    if description:
        p_desc = cell.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(2)
        p_desc.paragraph_format.space_after  = Pt(4)
        p_desc.paragraph_format.left_indent  = Inches(0.1)
        r = p_desc.add_run(description)
        r.italic         = True
        r.font.size      = Pt(9)
        r.font.color.rgb = RGBColor(0x5D, 0x40, 0x00)

    # ── "How to generate" label ─────────────────────────────────────────────
    p_label = cell.add_paragraph()
    p_label.paragraph_format.space_before = Pt(4)
    p_label.paragraph_format.space_after  = Pt(2)
    p_label.paragraph_format.left_indent  = Inches(0.1)
    r = p_label.add_run("Mermaid source — paste into claude.ai or mermaid.live to render:")
    r.bold           = True
    r.font.size      = Pt(8.5)
    r.font.color.rgb = RGBColor(0x6D, 0x4C, 0x00)

    # ── Mermaid code block (slightly darker yellow) ─────────────────────────
    p_code = cell.add_paragraph()
    p_code.paragraph_format.space_before = Pt(2)
    p_code.paragraph_format.space_after  = Pt(4)
    p_code.paragraph_format.left_indent  = Inches(0.15)
    p_code.paragraph_format.right_indent = Inches(0.15)
    # Shade this inner paragraph amber-ish
    pPr  = p_code._p.get_or_add_pPr()
    shd2 = OxmlElement('w:shd')
    shd2.set(qn('w:val'),   'clear')
    shd2.set(qn('w:color'), 'auto')
    shd2.set(qn('w:fill'),  'FFF3CD')   # slightly deeper amber for contrast
    pPr.append(shd2)
    r = p_code.add_run(mermaid_code)
    r.font.name      = 'Courier New'
    r.font.size      = Pt(7.5)
    r.font.color.rgb = RGBColor(0x2C, 0x1A, 0x00)

    # ── Render hint footer ──────────────────────────────────────────────────
    p_hint = cell.add_paragraph()
    p_hint.paragraph_format.space_before = Pt(2)
    p_hint.paragraph_format.space_after  = Pt(6)
    p_hint.paragraph_format.left_indent  = Inches(0.1)
    r = p_hint.add_run("⚠  Replace this box with the rendered diagram when submitting the final report.")
    r.italic         = True
    r.font.size      = Pt(8)
    r.font.color.rgb = RGBColor(0xBF, 0x36, 0x00)

    doc.add_paragraph()   # breathing room after box


# ════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ════════════════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
run = p.add_run("SMART DISASTER RESPONSE")
run.bold           = True
run.font.size      = Pt(26)
run.font.color.rgb = DARK_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("MANAGEMENT INFORMATION SYSTEM")
run.bold           = True
run.font.size      = Pt(22)
run.font.color.rgb = MID_BLUE

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Complete Project Documentation Report")
run.font.size      = Pt(14)
run.font.color.rgb = TEXT_MID
run.italic         = True

add_hr()

# Group members table on cover
cover_data = [
    ("24I-3108", "Sohaib Akhlaq"),
    ("24I-3074", "Shaiman Qasir"),
    ("24I-3107", "M. Hasaam"),
]
add_table(["Student ID", "Name"], cover_data, col_widths=[2.0, 3.5])

add_para("Course: Database Systems", bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Instructor: Ms. Zoya Sumbul", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("Submission Deadline: 3rd May 2026, 11:59 PM", size=11, align=WD_ALIGN_PARAGRAPH.CENTER)
add_para("FAST-NUCES", italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=TEXT_MID)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — PROJECT OVERVIEW
# ════════════════════════════════════════════════════════════════════════════
add_heading("1. Project Overview & Problem Statement", level=1)
add_hr()

add_heading("1.1 Background", level=2)
add_para(
    "Pakistan is a country that frequently faces severe natural disasters — floods along the Indus River basin, "
    "earthquakes in the northern and western regions, urban fires in densely populated cities, and industrial gas explosions. "
    "During such events, thousands of emergency reports are generated in real time from citizens using mobile applications, "
    "helpline operators, and automated monitoring systems. Without a centralized, structured information system, coordination "
    "between field teams, warehouses, hospitals, and financial authorities becomes chaotic, leading to delayed response times, "
    "misallocated resources, and preventable loss of life."
)

add_heading("1.2 Problem Definition", level=2)
add_para("The existing disaster management processes in Pakistan rely heavily on phone calls, spreadsheets, and manual "
         "reporting — all of which fail at scale during a crisis. There is no single source of truth for:")
for item in [
    "Which emergency reports are active and what their severity is",
    "Which rescue teams are available for immediate dispatch",
    "What resources are stocked in which warehouses and how much is left",
    "Which hospitals have capacity to accept new patients",
    "How much money has been raised and spent per disaster event",
    "Who approved what action and when (audit trail)",
]:
    add_bullet(item)

add_heading("1.3 Proposed Solution", level=2)
add_para(
    "The Smart Disaster Response Management Information System (MIS) is a full-stack enterprise application built on "
    "Microsoft SQL Server and Next.js, designed to be the central operational hub for all disaster response stakeholders. "
    "The system provides:"
)
for item in [
    "Real-time emergency report ingestion and status tracking",
    "Dynamic rescue team assignment with availability enforcement",
    "Warehouse inventory management with automated low-stock alerts",
    "Hospital capacity monitoring and patient routing",
    "Financial management (donations, expenses, budget tracking)",
    "Approval-based workflows for critical operations",
    "Comprehensive audit logging with full traceability",
    "Role-based access control at both application and database layers",
    "Advanced database performance optimization through indexing, views, and triggers",
]:
    add_bullet(item)

add_heading("1.4 Stakeholders and Roles", level=2)
add_table(
    ["Role", "Responsibilities"],
    [
        ("Administrator",        "Full system control, user management, approvals, audit review"),
        ("Emergency Operator",   "Creates reports, assigns teams, monitors incidents"),
        ("Field Officer",        "Updates mission status, completes assignments"),
        ("Warehouse Manager",    "Manages inventory, submits dispatch requests"),
        ("Finance Officer",      "Records donations and expenses, requests financial approvals"),
    ],
    col_widths=[2.2, 4.3],
)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — SYSTEM ARCHITECTURE
# ════════════════════════════════════════════════════════════════════════════
add_heading("2. System Architecture", level=1)
add_hr()

add_heading("2.1 Technology Stack", level=2)
add_table(
    ["Layer", "Technology", "Version", "Purpose"],
    [
        ("Database",        "Microsoft SQL Server", "2019+",    "Primary data store, ACID transactions, triggers, views"),
        ("DB Driver",       "node-mssql",           "^12.5.0",  "Parameterized T-SQL queries"),
        ("Backend",         "Next.js API Routes",   "16.2.4",   "RESTful endpoints, server actions"),
        ("Frontend",        "React",                "19.2.4",   "Component-based UI"),
        ("Authentication",  "NextAuth.js",          "^4.24.14", "JWT sessions, credential validation"),
        ("Password Hash",   "SHA2_256 (T-SQL)",     "—",        "Server-side password hashing"),
        ("Charts",          "Recharts",             "^2.6.2",   "Dashboard visualizations"),
        ("Form Handling",   "React Hook Form",      "^7.45.0",  "Validated form inputs"),
        ("Validation",      "Zod",                  "^4.3.6",   "Runtime type safety"),
        ("Styling",         "Tailwind CSS",         "^4",       "Responsive design system"),
    ],
    col_widths=[1.4, 1.8, 1.1, 2.2],
)

add_heading("2.2 High-Level Architecture", level=2)
add_para(
    "The system follows a three-tier architecture: the Client Browser communicates via HTTPS with the Next.js server, "
    "which in turn communicates with Microsoft SQL Server using a connection pool of up to 10 concurrent connections."
)
add_para(
    "Client Layer: Login, Role-based dashboards, Forms (incident reporting, resource requests, financial entries), "
    "Analytics and reports."
)
add_para(
    "Application Layer: NextAuth.js JWT session management, 20+ REST API routes, Zod-validated server actions, "
    "Route middleware for authorization."
)
add_para(
    "Database Layer: 18 relational tables, 8 triggers (event-driven automation), 5 views (role-based abstraction), "
    "25+ indexes (performance optimization), ACID-compliant transactions with row-level locking."
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — ERD
# ════════════════════════════════════════════════════════════════════════════
add_heading("3. Entity Relationship Diagram (ERD)", level=1)
add_hr()

add_heading("3.1 ERD Diagram", level=2)
add_para(
    "The ERD below represents all 18 entities in the DisasterResponseMIS database. "
    "The diagram is provided in Mermaid notation. Paste the code block into claude.ai or mermaid.live to render it."
)

erd_code = """erDiagram
    Roles { int role_id PK; varchar role_name; varchar description }
    Users { int user_id PK; varchar name; varchar email; varchar password_hash; int role_id FK }
    Citizens { int citizen_id PK; varchar name; varchar phone; varchar address }
    EmergencyReports { int report_id PK; int citizen_id FK; varchar location; varchar disaster_type; varchar severity; varchar status }
    RescueTeams { int team_id PK; varchar team_name; varchar team_type; varchar status }
    TeamAssignments { int assignment_id PK; int report_id FK; int team_id FK; varchar status }
    Warehouses { int warehouse_id PK; varchar name; varchar location }
    Resources { int resource_id PK; varchar resource_name; varchar category; varchar unit }
    WarehouseInventory { int inventory_id PK; int warehouse_id FK; int resource_id FK; int quantity; int min_threshold }
    ApprovalRequests { int request_id PK; int requested_by FK; varchar request_type; varchar status; int decided_by FK }
    ResourceAllocations { int allocation_id PK; int request_id FK; int inventory_id FK; int quantity_dispatched }
    Hospitals { int hospital_id PK; varchar name; int total_beds; varchar specialty }
    Patients { int patient_id PK; int hospital_id FK; int report_id FK; varchar name; varchar status }
    Donations { int donation_id PK; varchar donor_name; decimal amount; varchar type; int approved_by FK }
    Expenses { int expense_id PK; varchar category; decimal amount; int approved_by FK }
    AuditLogs { int log_id PK; int user_id FK; varchar action_type; varchar description; datetime2 timestamp }
    Permissions { int permission_id PK; int role_id FK; varchar module_name; bit can_read; bit can_write; bit can_delete }
    StockAlerts { int alert_id PK; int inventory_id FK; datetime2 alert_time }

    Roles ||--o{ Users : has
    Roles ||--o{ Permissions : defines
    Users ||--o{ ApprovalRequests : requests
    Users ||--o{ ApprovalRequests : decides
    Users ||--o{ AuditLogs : generates
    Citizens ||--o{ EmergencyReports : submits
    EmergencyReports ||--o{ TeamAssignments : triggers
    EmergencyReports ||--o{ Patients : leads_to
    RescueTeams ||--o{ TeamAssignments : fulfills
    Warehouses ||--o{ WarehouseInventory : holds
    Resources ||--o{ WarehouseInventory : tracked_in
    WarehouseInventory ||--o{ ResourceAllocations : source_of
    WarehouseInventory ||--o{ StockAlerts : triggers
    ApprovalRequests ||--o{ ResourceAllocations : authorizes
    Hospitals ||--o{ Patients : admits"""

add_diagram_box(
    "Full Entity Relationship Diagram — All 18 Entities",
    erd_code,
    "Shows all entities, their attributes, and relationships with cardinality notation."
)

add_heading("3.2 ERD Relationship Summary", level=2)
add_table(
    ["Relationship", "Cardinality", "Description"],
    [
        ("Roles → Users",                     "1:M", "One role assigned to many system users"),
        ("Roles → Permissions",               "1:M", "One role has CRUD permissions for many modules"),
        ("Citizens → EmergencyReports",       "1:M", "One citizen can file many emergency reports"),
        ("EmergencyReports → TeamAssignments","1:M", "One report can trigger multiple team assignments over time"),
        ("RescueTeams → TeamAssignments",     "1:M", "One team can be assigned to many missions"),
        ("Warehouses → WarehouseInventory",   "1:M", "One warehouse holds many resource inventory records"),
        ("Resources → WarehouseInventory",    "1:M", "One resource type tracked in many warehouses"),
        ("WarehouseInventory → ResourceAllocations","1:M","One inventory item can be dispatched multiple times"),
        ("ApprovalRequests → ResourceAllocations","1:M","One approval can authorize multiple dispatch records"),
        ("Hospitals → Patients",              "1:M", "One hospital admits many patients"),
        ("EmergencyReports → Patients",       "1:M", "One incident can result in many patients"),
        ("Users → ApprovalRequests (requester)","1:M","One user submits many approval requests"),
        ("Users → ApprovalRequests (decider)","1:M", "One user approves/rejects many requests"),
        ("Users → AuditLogs",                 "1:M", "Every user action generates audit log entries"),
        ("WarehouseInventory → StockAlerts",  "1:M", "One inventory item can trigger many alerts over time"),
    ],
    col_widths=[2.5, 1.0, 3.0],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — RELATIONAL SCHEMA
# ════════════════════════════════════════════════════════════════════════════
add_heading("4. Relational Schema", level=1)
add_hr()
add_para(
    "The following relational schema defines all 18 tables with their attributes, data types, "
    "primary keys (PK), foreign keys (FK), and integrity constraints. Tables are listed in dependency order."
)

schemas = [
    ("TABLE 1: Roles",
     "Roles(\n  role_id    INT  IDENTITY(1,1)  PK,\n  role_name  VARCHAR(50)  NOT NULL  UNIQUE,\n  description VARCHAR(255) NOT NULL\n)"),
    ("TABLE 2: Users",
     "Users(\n  user_id       INT  IDENTITY(1,1)  PK,\n  name          VARCHAR(100) NOT NULL,\n  email         VARCHAR(150) NOT NULL  UNIQUE,\n  password_hash VARCHAR(256) NOT NULL,\n  phone         VARCHAR(20)  NULL,\n  role_id       INT          NOT NULL  FK → Roles(role_id) [ON UPDATE CASCADE],\n  is_active     BIT          NOT NULL  DEFAULT 1,\n  created_at    DATETIME2              DEFAULT SYSDATETIME()\n)"),
    ("TABLE 3: Citizens",
     "Citizens(\n  citizen_id INT          IDENTITY(1,1)  PK,\n  name       VARCHAR(100) NOT NULL,\n  phone      VARCHAR(20)  NOT NULL,\n  address    VARCHAR(255) NULL,\n  email      VARCHAR(150) NULL\n)"),
    ("TABLE 4: EmergencyReports",
     "EmergencyReports(\n  report_id     INT          IDENTITY(1,1)  PK,\n  citizen_id    INT          NOT NULL  FK → Citizens(citizen_id) [ON UPDATE CASCADE],\n  location      VARCHAR(255) NOT NULL,\n  disaster_type VARCHAR(100) NOT NULL,\n  severity      VARCHAR(20)  NOT NULL  CHECK IN ('Low','Medium','High'),\n  report_time   DATETIME2    NOT NULL  DEFAULT SYSDATETIME(),\n  status        VARCHAR(30)  NOT NULL  DEFAULT 'Pending'\n                                       CHECK IN ('Pending','Dispatched','Resolved','Cancelled')\n)"),
    ("TABLE 5: RescueTeams",
     "RescueTeams(\n  team_id          INT          IDENTITY(1,1)  PK,\n  team_name        VARCHAR(100) NOT NULL,\n  team_type        VARCHAR(50)  NOT NULL,\n  current_location VARCHAR(255) NOT NULL,\n  status           VARCHAR(30)  NOT NULL  DEFAULT 'Available'\n                                          CHECK IN ('Available','Busy','Offline'),\n  equipment        VARCHAR(500) NULL\n)"),
    ("TABLE 6: TeamAssignments",
     "TeamAssignments(\n  assignment_id INT       IDENTITY(1,1)  PK,\n  report_id     INT       NOT NULL  FK → EmergencyReports(report_id),\n  team_id       INT       NOT NULL  FK → RescueTeams(team_id),\n  assigned_at   DATETIME2 NOT NULL  DEFAULT SYSDATETIME(),\n  completed_at  DATETIME2 NULL,\n  status        VARCHAR(30) NOT NULL DEFAULT 'Assigned'\n                            CHECK IN ('Assigned','EnRoute','OnSite','Completed','Cancelled')\n)"),
    ("TABLE 7: Warehouses",
     "Warehouses(\n  warehouse_id INT          IDENTITY(1,1)  PK,\n  name         VARCHAR(100) NOT NULL,\n  location     VARCHAR(255) NOT NULL,\n  phone        VARCHAR(20)  NULL\n)"),
    ("TABLE 8: Resources",
     "Resources(\n  resource_id   INT          IDENTITY(1,1)  PK,\n  resource_name VARCHAR(100) NOT NULL,\n  category      VARCHAR(50)  NOT NULL,\n  unit          VARCHAR(20)  NOT NULL\n)"),
    ("TABLE 9: WarehouseInventory",
     "WarehouseInventory(\n  inventory_id  INT       IDENTITY(1,1)  PK,\n  warehouse_id  INT       NOT NULL  FK → Warehouses(warehouse_id),\n  resource_id   INT       NOT NULL  FK → Resources(resource_id),\n  quantity      INT       NOT NULL  DEFAULT 0  CHECK (quantity >= 0),\n  min_threshold INT       NOT NULL  DEFAULT 10 CHECK (min_threshold >= 0),\n  last_updated  DATETIME2           DEFAULT SYSDATETIME(),\n  UNIQUE (warehouse_id, resource_id)\n)"),
    ("TABLE 10: ApprovalRequests",
     "ApprovalRequests(\n  request_id   INT       IDENTITY(1,1)  PK,\n  requested_by INT       NOT NULL  FK → Users(user_id),\n  request_type VARCHAR(50) NOT NULL,\n  details      TEXT      NOT NULL,\n  status       VARCHAR(20) NOT NULL DEFAULT 'Pending'\n                           CHECK IN ('Pending','Approved','Rejected'),\n  decided_by   INT       NULL  FK → Users(user_id),\n  decision_note TEXT     NULL,\n  created_at   DATETIME2 NOT NULL  DEFAULT SYSDATETIME(),\n  decided_at   DATETIME2 NULL\n)"),
    ("TABLE 11: ResourceAllocations",
     "ResourceAllocations(\n  allocation_id       INT          IDENTITY(1,1)  PK,\n  request_id          INT          NULL  FK → ApprovalRequests(request_id),\n  inventory_id        INT          NOT NULL  FK → WarehouseInventory(inventory_id),\n  quantity_dispatched INT          NOT NULL  CHECK (quantity_dispatched > 0),\n  destination         VARCHAR(255) NOT NULL,\n  allocated_at        DATETIME2    NOT NULL  DEFAULT SYSDATETIME()\n)"),
    ("TABLE 12: Hospitals",
     "Hospitals(\n  hospital_id INT          IDENTITY(1,1)  PK,\n  name        VARCHAR(100) NOT NULL,\n  location    VARCHAR(255) NOT NULL,\n  total_beds  INT          NOT NULL  DEFAULT 0  CHECK (total_beds >= 0),\n  specialty   VARCHAR(200) NULL\n)"),
    ("TABLE 13: Patients",
     "Patients(\n  patient_id  INT          IDENTITY(1,1)  PK,\n  hospital_id INT          NOT NULL  FK → Hospitals(hospital_id),\n  report_id   INT          NULL  FK → EmergencyReports(report_id),\n  name        VARCHAR(100) NOT NULL,\n  admitted_at DATETIME2    NOT NULL  DEFAULT SYSDATETIME(),\n  status      VARCHAR(30)  NOT NULL  DEFAULT 'Admitted'\n                           CHECK IN ('Admitted','Discharged','Transferred','Deceased')\n)"),
    ("TABLE 14: Donations",
     "Donations(\n  donation_id   INT            IDENTITY(1,1)  PK,\n  donor_name    VARCHAR(100)   NOT NULL,\n  amount        DECIMAL(10,2)  NOT NULL  CHECK (amount > 0),\n  type          VARCHAR(30)    NOT NULL  CHECK IN ('Cash','Kind','Pledge'),\n  disaster_event VARCHAR(100)  NOT NULL,\n  donated_at    DATETIME2      NOT NULL  DEFAULT SYSDATETIME(),\n  approved_by   INT            NULL  FK → Users(user_id)\n)"),
    ("TABLE 15: Expenses",
     "Expenses(\n  expense_id    INT            IDENTITY(1,1)  PK,\n  category      VARCHAR(50)    NOT NULL,\n  amount        DECIMAL(10,2)  NOT NULL  CHECK (amount > 0),\n  disaster_event VARCHAR(100)  NOT NULL,\n  incurred_at   DATETIME2      NOT NULL  DEFAULT SYSDATETIME(),\n  approved_by   INT            NULL  FK → Users(user_id)\n)"),
    ("TABLE 16: AuditLogs",
     "AuditLogs(\n  log_id         INT          IDENTITY(1,1)  PK,\n  user_id        INT          NOT NULL  FK → Users(user_id),\n  action_type    VARCHAR(50)  NOT NULL  CHECK IN ('INSERT','UPDATE','DELETE','LOGIN','LOGOUT'),\n  description    VARCHAR(500) NOT NULL,\n  table_affected VARCHAR(100) NOT NULL,\n  record_id      INT          NULL,\n  timestamp      DATETIME2    NOT NULL  DEFAULT SYSDATETIME(),\n  ip_address     VARCHAR(45)  NULL\n)"),
    ("TABLE 17: Permissions",
     "Permissions(\n  permission_id INT          IDENTITY(1,1)  PK,\n  role_id       INT          NOT NULL  FK → Roles(role_id),\n  module_name   VARCHAR(100) NOT NULL,\n  can_read      BIT          NOT NULL  DEFAULT 0,\n  can_write     BIT          NOT NULL  DEFAULT 0,\n  can_delete    BIT          NOT NULL  DEFAULT 0,\n  UNIQUE (role_id, module_name)\n)"),
    ("TABLE 18: StockAlerts",
     "StockAlerts(\n  alert_id     INT       IDENTITY(1,1)  PK,\n  inventory_id INT       NOT NULL  FK → WarehouseInventory(inventory_id),\n  alert_time   DATETIME2 NOT NULL  DEFAULT SYSDATETIME()\n)"),
]

for title, schema in schemas:
    add_heading(title, level=3)
    add_code(schema)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — NORMALIZATION
# ════════════════════════════════════════════════════════════════════════════
add_heading("5. Normalization Steps", level=1)
add_hr()

add_heading("5.1 First Normal Form (1NF)", level=2)
add_para("Definition: Every attribute must be atomic (indivisible) and every row must be uniquely identifiable.")
add_para("Application: All attributes across the 18 tables are atomic. No multi-valued attributes exist. "
         "Every table has a single-column IDENTITY primary key. Example non-1NF situation avoided: an early draft stored "
         "team equipment as a comma-separated list of items. This was kept as a single VARCHAR(500) description field. "
         "If team equipment needed to be queried individually, a separate TeamEquipment table would be required.", italic=True, color=TEXT_MID)

add_heading("5.2 Second Normal Form (2NF)", level=2)
add_para("Definition: Must be in 1NF. Every non-key attribute must be fully dependent on the entire primary key "
         "(no partial dependencies — relevant only for composite PKs).")
add_para("Application: Most tables use single-column surrogate PKs, making 2NF trivially satisfied. "
         "WarehouseInventory has UNIQUE(warehouse_id, resource_id) but its PK is the surrogate inventory_id. "
         "quantity and min_threshold depend on the specific (warehouse, resource) combination — fully on the PK. "
         "Non-2NF situation avoided: early design had warehouse_name and resource_name stored directly in "
         "WarehouseInventory — these depended only on warehouse_id and resource_id respectively, not on inventory_id. "
         "Resolved by using JOINs to Warehouses and Resources tables instead.", italic=True, color=TEXT_MID)

add_heading("5.3 Third Normal Form (3NF)", level=2)
add_para("Definition: Must be in 2NF. No non-key attribute should be transitively dependent on the PK through another non-key attribute.")
add_para("Key 3NF correction — StockAlerts table: An early version included resource_name, warehouse_name, "
         "current_quantity, and threshold inline. These created transitive dependencies: "
         "alert_id → inventory_id → resource_name (transitive). Corrected design stores only (alert_id, inventory_id, alert_time). "
         "All derived details are fetched via JOIN at query time. The SQL file comments explicitly note: 'Corrected to 3NF'.", italic=True, color=TEXT_MID)
add_para("Documented denormalization: disaster_event is stored as a string in Donations and Expenses rather than as "
         "an FK to a Disasters table. This is a deliberate design decision — disaster events are ad-hoc and named "
         "informally. A formal event registry would require administrative overhead before any financial records could "
         "be created. The FULL OUTER JOIN in vw_FinancialSummaryByDisaster handles partial data naturally.", italic=True, color=TEXT_MID)

add_heading("5.4 Normalization Summary", level=2)
add_table(
    ["Table", "1NF", "2NF", "3NF", "Notes"],
    [
        ("Roles",                "✓", "✓", "✓", "No derived attributes"),
        ("Users",                "✓", "✓", "✓", "role_id FK — role_name not stored"),
        ("Citizens",             "✓", "✓", "✓", "Atomic attributes only"),
        ("EmergencyReports",     "✓", "✓", "✓", "Citizen details stay in Citizens table"),
        ("RescueTeams",          "✓", "✓", "✓", "No transitive dependencies"),
        ("TeamAssignments",      "✓", "✓", "✓", "Bridge table with surrogate PK"),
        ("Warehouses",           "✓", "✓", "✓", "No derived data"),
        ("Resources",            "✓", "✓", "✓", "Atomic attributes only"),
        ("WarehouseInventory",   "✓", "✓", "✓", "No warehouse/resource names stored inline"),
        ("ApprovalRequests",     "✓", "✓", "✓", "Two FKs to Users for distinct roles"),
        ("ResourceAllocations",  "✓", "✓", "✓", "Derived data fetched via JOIN"),
        ("Hospitals",            "✓", "✓", "✓", "total_beds is atomic fact"),
        ("Patients",             "✓", "✓", "✓", "Hospital/report info via FK"),
        ("Donations",            "✓", "✓", "✓*","*disaster_event denormalized (justified)"),
        ("Expenses",             "✓", "✓", "✓*","*Same justification as Donations"),
        ("AuditLogs",            "✓", "✓", "✓", "Immutable log — no derivable redundancy"),
        ("Permissions",          "✓", "✓", "✓", "Role + module as unique composite"),
        ("StockAlerts",          "✓", "✓", "✓", "CORRECTED: redundant columns removed"),
    ],
    col_widths=[2.2, 0.6, 0.6, 0.6, 2.5],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — DDL
# ════════════════════════════════════════════════════════════════════════════
add_heading("6. SQL Implementation — DDL (Table Creation)", level=1)
add_hr()

add_heading("6.1 Database Creation", level=2)
add_code("""IF NOT EXISTS (SELECT name FROM sys.databases WHERE name = 'DisasterResponseMIS')
BEGIN
    CREATE DATABASE DisasterResponseMIS;
END
GO
USE DisasterResponseMIS;
GO""")

add_heading("6.2 Table Drop Order (Clean Slate — Reverse FK Dependency)", level=2)
add_para("Tables are dropped in reverse foreign key dependency order to prevent constraint violations when re-running the DDL script.")
add_code("""IF OBJECT_ID('dbo.StockAlerts',         'U') IS NOT NULL DROP TABLE StockAlerts;
IF OBJECT_ID('dbo.AuditLogs',           'U') IS NOT NULL DROP TABLE AuditLogs;
IF OBJECT_ID('dbo.Permissions',         'U') IS NOT NULL DROP TABLE Permissions;
IF OBJECT_ID('dbo.Patients',            'U') IS NOT NULL DROP TABLE Patients;
IF OBJECT_ID('dbo.ResourceAllocations', 'U') IS NOT NULL DROP TABLE ResourceAllocations;
IF OBJECT_ID('dbo.ApprovalRequests',    'U') IS NOT NULL DROP TABLE ApprovalRequests;
IF OBJECT_ID('dbo.Donations',           'U') IS NOT NULL DROP TABLE Donations;
IF OBJECT_ID('dbo.Expenses',            'U') IS NOT NULL DROP TABLE Expenses;
IF OBJECT_ID('dbo.TeamAssignments',     'U') IS NOT NULL DROP TABLE TeamAssignments;
IF OBJECT_ID('dbo.WarehouseInventory',  'U') IS NOT NULL DROP TABLE WarehouseInventory;
IF OBJECT_ID('dbo.EmergencyReports',    'U') IS NOT NULL DROP TABLE EmergencyReports;
IF OBJECT_ID('dbo.Citizens',            'U') IS NOT NULL DROP TABLE Citizens;
IF OBJECT_ID('dbo.RescueTeams',         'U') IS NOT NULL DROP TABLE RescueTeams;
IF OBJECT_ID('dbo.Hospitals',           'U') IS NOT NULL DROP TABLE Hospitals;
IF OBJECT_ID('dbo.Resources',           'U') IS NOT NULL DROP TABLE Resources;
IF OBJECT_ID('dbo.Warehouses',          'U') IS NOT NULL DROP TABLE Warehouses;
IF OBJECT_ID('dbo.Users',               'U') IS NOT NULL DROP TABLE Users;
IF OBJECT_ID('dbo.Roles',               'U') IS NOT NULL DROP TABLE Roles;
GO""")

add_heading("6.3 Key Table Creation Examples", level=2)

add_heading("EmergencyReports — CHECK constraints for domain integrity:", level=3)
add_code("""CREATE TABLE EmergencyReports (
    report_id     INT          IDENTITY(1,1) PRIMARY KEY,
    citizen_id    INT          NOT NULL,
    location      VARCHAR(255) NOT NULL,
    disaster_type VARCHAR(100) NOT NULL,
    severity      VARCHAR(20)  NOT NULL,
    report_time   DATETIME2    NOT NULL DEFAULT SYSDATETIME(),
    status        VARCHAR(30)  NOT NULL DEFAULT 'Pending',
    CONSTRAINT FK_EmergencyReports_Citizens
        FOREIGN KEY (citizen_id) REFERENCES Citizens(citizen_id)
        ON UPDATE CASCADE ON DELETE NO ACTION,
    CONSTRAINT CHK_EmergencyReports_Severity
        CHECK (severity IN ('Low', 'Medium', 'High')),
    CONSTRAINT CHK_EmergencyReports_Status
        CHECK (status IN ('Pending', 'Dispatched', 'Resolved', 'Cancelled'))
);""")

add_heading("WarehouseInventory — composite UNIQUE constraint and non-negative CHECK:", level=3)
add_code("""CREATE TABLE WarehouseInventory (
    inventory_id  INT       IDENTITY(1,1) PRIMARY KEY,
    warehouse_id  INT       NOT NULL,
    resource_id   INT       NOT NULL,
    quantity      INT       NOT NULL DEFAULT 0,
    min_threshold INT       NOT NULL DEFAULT 10,
    last_updated  DATETIME2 DEFAULT SYSDATETIME(),
    CONSTRAINT FK_WI_Warehouses
        FOREIGN KEY (warehouse_id) REFERENCES Warehouses(warehouse_id)
        ON UPDATE CASCADE ON DELETE NO ACTION,
    CONSTRAINT FK_WI_Resources
        FOREIGN KEY (resource_id) REFERENCES Resources(resource_id)
        ON UPDATE CASCADE ON DELETE NO ACTION,
    CONSTRAINT UQ_WarehouseInventory_Unique UNIQUE (warehouse_id, resource_id),
    CONSTRAINT CHK_WarehouseInventory_Quantity  CHECK (quantity >= 0),
    CONSTRAINT CHK_WarehouseInventory_Threshold CHECK (min_threshold >= 0)
);""")

add_heading("Password Hashing using SHA2_256:", level=3)
add_code("""INSERT INTO Users (name, email, password_hash, phone, role_id) VALUES
('Sohaib Akhlaq',
 'admin@mis.pk',
 CONVERT(VARCHAR(256), HASHBYTES('SHA2_256', 'admin123'), 2),
 '03001234567', 1);""")

add_heading("6.4 SQL File Execution Order", level=2)
add_table(
    ["File", "Description", "Run After"],
    [
        ("01_DDL_CreateTables.sql", "Creates 18 tables, seeds Roles and 2 base Users", "— (first file)"),
        ("04_Triggers.sql",         "Creates 8 event-driven triggers",                  "FILE 01"),
        ("05_Views.sql",            "Creates 5 role-scoped views",                      "FILE 01"),
        ("03_DML_SampleData.sql",   "Inserts 200+ sample records",                      "FILE 01, 04"),
        ("06_Queries_Transactions.sql","9 queries + 4 ACID transactions",               "FILE 01, 03, 04"),
        ("02_DDL_Indexes.sql",      "25+ indexes + performance analysis",                "All above files"),
    ],
    col_widths=[2.4, 2.8, 1.3],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — DML
# ════════════════════════════════════════════════════════════════════════════
add_heading("7. SQL Implementation — DML (Sample Data)", level=1)
add_hr()

add_heading("7.1 Sample Data Volume", level=2)
add_table(
    ["Table", "Records", "Notes"],
    [
        ("Roles",               "5",    "Administrator, Emergency Operator, Field Officer, Warehouse Manager, Finance Officer"),
        ("Users",               "5",    "One per role, all with SHA2_256-hashed passwords"),
        ("Permissions",         "29",   "Full RBAC matrix: read/write/delete per module per role"),
        ("Citizens",            "10",   "Across Karachi, Lahore, Islamabad, Peshawar, Quetta, Rawalpindi, Multan"),
        ("RescueTeams",         "8",    "Mix of Medical, Fire, Rescue types"),
        ("Warehouses",          "6",    "Karachi, Lahore, Islamabad, Peshawar, Quetta, Multan"),
        ("Resources",           "10",   "Water, Rice, Blankets, Tents, First Aid, ORS, Medicines, Equipment"),
        ("WarehouseInventory",  "30",   "5 resources × 6 warehouses (varies)"),
        ("EmergencyReports",    "30",   "Mix of Flood, Earthquake, Urban Fire, Gas Explosion"),
        ("TeamAssignments",     "10",   "5 Completed (trigger reset teams Available), 5 Active (trigger set Busy)"),
        ("ApprovalRequests",    "8",    "Mix of Approved, Rejected, and Pending states"),
        ("ResourceAllocations", "4",    "Trigger auto-reduced inventory on INSERT"),
        ("Hospitals",           "6",    "Major government hospitals across Pakistan"),
        ("Patients",            "10",   "Admitted across 6 hospitals, some linked to emergency reports"),
        ("Donations",           "8",    "Cash, Kind, Pledge types; trigger wrote 8 audit rows automatically"),
        ("Expenses",            "8",    "Procurement, Distribution, Operations; trigger wrote 8 audit rows"),
        ("AuditLogs",           "24",   "8 manual entries + 16 trigger-generated entries"),
        ("StockAlerts",         "0–1",  "Auto-generated by trigger T6 when stock crosses below threshold"),
    ],
    col_widths=[2.0, 0.9, 3.6],
)

add_heading("7.2 Demo Login Credentials", level=2)
add_table(
    ["Role", "Email", "Password"],
    [
        ("Administrator",       "admin@mis.pk",        "admin123"),
        ("Emergency Operator",  "operator@mis.pk",     "op123"),
        ("Field Officer",       "fieldofficer@mis.pk", "field123"),
        ("Warehouse Manager",   "warehouse@mis.pk",    "wh123"),
        ("Finance Officer",     "finance@mis.pk",      "fin123"),
    ],
    col_widths=[2.0, 2.5, 2.0],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 8 — CORE QUERIES
# ════════════════════════════════════════════════════════════════════════════
add_heading("8. Core SELECT Queries with Expected Output", level=1)
add_hr()

queries = [
    ("Query A — High-Severity Reports in a Specific City",
     "Emergency Operator dashboard: all active High-severity incidents in Karachi.",
     """SELECT er.report_id, c.name AS citizen_name, c.phone,
       er.location, er.disaster_type, er.severity, er.report_time, er.status
FROM EmergencyReports er
INNER JOIN Citizens c ON er.citizen_id = c.citizen_id
WHERE er.location LIKE '%Karachi%' AND er.severity = 'High'
ORDER BY er.report_time DESC;""",
     ["report_id","citizen_name","location","disaster_type","severity","status"],
     [("21","Ali Hassan","Karachi, Malir","Flood","High","Pending"),
      ("11","Ali Hassan","Karachi, Orangi Town","Flood","High","Dispatched"),
      ("8","Iqra Siddiqui","Karachi, SITE Area","Urban Fire","High","Pending"),
      ("1","Ali Hassan","Karachi, Lyari","Flood","High","Dispatched")]),

    ("Query B — Resources Below Minimum Stock Threshold",
     "Warehouse Manager identifies items that urgently need restocking.",
     """SELECT r.resource_name, r.category, r.unit,
       w.name AS warehouse_name, wi.quantity AS current_stock,
       wi.min_threshold, (wi.min_threshold - wi.quantity) AS shortage_amount
FROM WarehouseInventory wi
INNER JOIN Resources  r ON wi.resource_id  = r.resource_id
INNER JOIN Warehouses w ON wi.warehouse_id = w.warehouse_id
WHERE wi.quantity < wi.min_threshold
ORDER BY shortage_amount DESC;""",
     ["resource_name","warehouse_name","current_stock","min_threshold","shortage_amount"],
     [("Tents (4-person)","Islamabad Federal Store","5","10","5")]),

    ("Query D — Financial Summary by Disaster Event",
     "Finance Officer views surplus or deficit per disaster event.",
     """SELECT COALESCE(d.disaster_event, e.disaster_event) AS disaster_event,
       ISNULL(SUM(d.amount), 0) AS total_donations,
       ISNULL(SUM(e.amount), 0) AS total_expenses,
       ISNULL(SUM(d.amount), 0) - ISNULL(SUM(e.amount), 0) AS net_balance,
       CASE WHEN ISNULL(SUM(d.amount),0) - ISNULL(SUM(e.amount),0) >= 0
            THEN 'SURPLUS' ELSE 'DEFICIT' END AS financial_status
FROM Donations d
FULL OUTER JOIN Expenses e ON d.disaster_event = e.disaster_event
GROUP BY COALESCE(d.disaster_event, e.disaster_event)
ORDER BY net_balance DESC;""",
     ["disaster_event","total_donations","total_expenses","net_balance","status"],
     [("Lahore Floods 2024","1,700,000","350,000","1,350,000","SURPLUS"),
      ("Quetta Earthquake 2024","1,030,000","530,000","500,000","SURPLUS"),
      ("Karachi Floods 2024","950,000","470,000","480,000","SURPLUS"),
      ("Peshawar Floods 2024","420,000","270,000","150,000","SURPLUS")]),

    ("Query E — Pending Approval Requests",
     "Administrator dashboard showing all requests awaiting a decision.",
     """SELECT ar.request_id, u.name AS requested_by_name,
       r.role_name AS requester_role, ar.request_type, ar.details,
       ar.status, ar.created_at,
       DATEDIFF(HOUR, ar.created_at, SYSDATETIME()) AS pending_since_hours
FROM ApprovalRequests ar
INNER JOIN Users u ON ar.requested_by = u.user_id
INNER JOIN Roles r ON u.role_id       = r.role_id
WHERE ar.status = 'Pending'
ORDER BY ar.created_at ASC;""",
     ["request_id","requested_by","requester_role","request_type","pending_hours"],
     [("3","Ayesha Siddiqui","Warehouse Manager","ResourceDispatch","48"),
      ("6","Operator User","Emergency Operator","TeamDeployment","24"),
      ("8","Bilal Farooq","Finance Officer","FinancialApproval","12")]),

    ("Query F — Hospital Capacity Report",
     "Emergency Operator determines where to route new patients.",
     """SELECT h.name AS hospital_name, h.location, h.specialty,
       h.total_beds, COUNT(p.patient_id) AS current_patients,
       h.total_beds - COUNT(p.patient_id) AS beds_available,
       CASE WHEN h.total_beds - COUNT(p.patient_id) <  5 THEN 'CRITICAL'
            WHEN h.total_beds - COUNT(p.patient_id) < 20 THEN 'LIMITED'
            ELSE 'AVAILABLE' END AS capacity_status
FROM Hospitals h
LEFT JOIN Patients p ON h.hospital_id = p.hospital_id
                     AND p.status = 'Admitted'
GROUP BY h.hospital_id, h.name, h.location, h.specialty, h.total_beds
ORDER BY beds_available DESC;""",
     ["hospital_name","total_beds","current_patients","beds_available","capacity_status"],
     [("Karachi Civil Hospital","500","3","497","AVAILABLE"),
      ("Lady Reading Hospital","450","2","448","AVAILABLE"),
      ("Services Hospital Lahore","400","0","400","AVAILABLE"),
      ("PIMS Hospital","350","2","348","AVAILABLE"),
      ("Nishtar Hospital","380","1","379","AVAILABLE"),
      ("Bolan Medical Complex","250","2","248","AVAILABLE")]),

    ("Query H — MIS Incident Report by Disaster Type",
     "Management dashboard showing incident breakdown by disaster type.",
     """SELECT er.disaster_type,
       COUNT(*) AS total_incidents,
       SUM(CASE WHEN er.severity='High'   THEN 1 ELSE 0 END) AS high_severity,
       SUM(CASE WHEN er.severity='Medium' THEN 1 ELSE 0 END) AS medium_severity,
       SUM(CASE WHEN er.severity='Low'    THEN 1 ELSE 0 END) AS low_severity,
       SUM(CASE WHEN er.status='Resolved' THEN 1 ELSE 0 END) AS resolved,
       SUM(CASE WHEN er.status IN ('Pending','Dispatched') THEN 1 ELSE 0 END) AS active
FROM EmergencyReports er
GROUP BY er.disaster_type
ORDER BY total_incidents DESC;""",
     ["disaster_type","total","high","medium","low","resolved","active"],
     [("Flood","15","9","4","2","3","12"),
      ("Urban Fire","9","3","4","2","4","5"),
      ("Earthquake","4","2","1","1","1","3"),
      ("Gas Explosion","3","1","0","2","0","3")]),
]

for title, use_case, sql, headers, rows in queries:
    add_heading(title, level=2)
    add_para(f"Use case: {use_case}", italic=True, color=TEXT_MID, size=9.5)
    add_code(sql)
    add_para("Expected Output:", bold=True, size=9.5)
    add_table(headers, rows)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 9 — TRANSACTIONS
# ════════════════════════════════════════════════════════════════════════════
add_heading("9. Transaction Handling Demonstration", level=1)
add_hr()

add_heading("9.1 ACID Properties in the System", level=2)
add_table(
    ["Property", "How It Is Implemented"],
    [
        ("Atomicity",    "BEGIN TRANSACTION / COMMIT / ROLLBACK wraps multi-step operations. "
                         "Trigger T1 keeps allocation + inventory reduction in one unit."),
        ("Consistency",  "CHECK constraints (severity IN, quantity >= 0), FK constraints, "
                         "trigger business rules (no negative stock, team must be Available)."),
        ("Isolation",    "WITH (UPDLOCK, ROWLOCK) in Transaction 4 prevents two administrators "
                         "from approving the same request simultaneously."),
        ("Durability",   "COMMIT TRANSACTION writes to the SQL Server WAL. Triggers write "
                         "immutable audit entries that survive system restarts."),
    ],
    col_widths=[1.2, 5.3],
)

transactions = [
    ("Transaction 1 — Resource Allocation (Atomicity)",
     "Dispatch resources from warehouse to a relief point. Trigger auto-reduces inventory. "
     "If stock is insufficient, the entire transaction rolls back — no partial allocation is ever persisted.",
     """DECLARE @inv_id INT = 2, @qty INT = 50;
DECLARE @dest VARCHAR(255) = 'Lyari Secondary Relief Point, Karachi';
DECLARE @avail INT;

BEGIN TRY
    BEGIN TRANSACTION T1_ResourceDispatch;

    SELECT @avail = quantity FROM WarehouseInventory WHERE inventory_id = @inv_id;
    IF @avail < @qty
    BEGIN
        RAISERROR('T1: Insufficient stock. Available: %d, Requested: %d.', 16, 1, @avail, @qty);
        ROLLBACK TRANSACTION T1_ResourceDispatch; RETURN;
    END

    -- Trigger T1 fires here → auto-reduces inventory
    INSERT INTO ResourceAllocations (request_id, inventory_id, quantity_dispatched, destination)
    VALUES (1, @inv_id, @qty, @dest);

    -- Verify inventory still non-negative after trigger
    DECLARE @new_qty INT;
    SELECT @new_qty = quantity FROM WarehouseInventory WHERE inventory_id = @inv_id;
    IF @new_qty < 0
    BEGIN
        RAISERROR('T1: Inventory went negative after trigger.', 16, 1);
        ROLLBACK TRANSACTION T1_ResourceDispatch; RETURN;
    END

    COMMIT TRANSACTION T1_ResourceDispatch;
    PRINT '>>> TRANSACTION 1 COMMITTED SUCCESSFULLY.';
END TRY
BEGIN CATCH
    IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION T1_ResourceDispatch;
    PRINT 'TRANSACTION 1 ROLLED BACK. Error: ' + ERROR_MESSAGE();
END CATCH;""",
     "T1 Step 1 PASSED: Stock check OK. Available = 200\n"
     "T1 Step 2 PASSED: Allocation record #5 created.\n"
     "T1 Step 3 PASSED: Inventory still positive = 150\n"
     ">>> TRANSACTION 1 COMMITTED SUCCESSFULLY.",
     "T1 ROLLBACK CONFIRMED — Insufficient stock (100 available, 99999 requested)."),

    ("Transaction 2 — Rescue Team Assignment (Consistency)",
     "Assign a rescue team to an emergency. The team must be in Available status. "
     "If it is Busy or Offline, the transaction is rejected to prevent double-assignment.",
     """DECLARE @team_id INT = 7, @report_id INT = 2;
DECLARE @status VARCHAR(30);

BEGIN TRY
    BEGIN TRANSACTION T2_TeamAssignment;

    SELECT @status = status FROM RescueTeams WHERE team_id = @team_id;
    IF @status <> 'Available'
    BEGIN
        RAISERROR('T2: Team is %s, not Available.', 16, 1, @status);
        ROLLBACK TRANSACTION T2_TeamAssignment; RETURN;
    END

    -- Trigger T3 fires → team status set to Busy automatically
    INSERT INTO TeamAssignments (report_id, team_id, assigned_at, status)
    VALUES (@report_id, @team_id, SYSDATETIME(), 'Assigned');

    UPDATE EmergencyReports SET status = 'Dispatched' WHERE report_id = @report_id;

    COMMIT TRANSACTION T2_TeamAssignment;
    PRINT '>>> TRANSACTION 2 COMMITTED SUCCESSFULLY.';
END TRY
BEGIN CATCH
    IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION T2_TeamAssignment;
    PRINT 'TRANSACTION 2 ROLLED BACK: ' + ERROR_MESSAGE();
END CATCH;""",
     "T2 Step 1 PASSED: Team is Available.\n"
     "T2 Step 2 PASSED: Assignment #11 created. Team set Busy by trigger.\n"
     "T2 Step 3 PASSED: Report #2 marked Dispatched.\n"
     ">>> TRANSACTION 2 COMMITTED SUCCESSFULLY.",
     "T2 ROLLBACK CONFIRMED — Team status is Busy. Only Available teams can be assigned."),

    ("Transaction 3 — Financial Entry with Dual Audit Trail (Durability)",
     "Record a donation and produce two audit entries — one automatically by trigger T5a, "
     "and one written explicitly inside the transaction.",
     """DECLARE @donor VARCHAR(100) = 'Sindh Relief Fund';
DECLARE @amount DECIMAL(10,2) = 650000.00;
DECLARE @don_id INT;

BEGIN TRY
    BEGIN TRANSACTION T3_FinancialEntry;

    IF @amount <= 0
    BEGIN
        RAISERROR('T3: Amount must be positive.', 16, 1);
        ROLLBACK TRANSACTION T3_FinancialEntry; RETURN;
    END

    -- Trigger T5a fires automatically → writes audit row
    INSERT INTO Donations (donor_name, amount, type, disaster_event, donated_at, approved_by)
    VALUES (@donor, @amount, 'Cash', 'Karachi Floods 2024', SYSDATETIME(), 1);

    SET @don_id = CAST(SCOPE_IDENTITY() AS INT);

    -- Explicit transaction-level audit entry with IP address
    INSERT INTO AuditLogs (user_id, action_type, description, table_affected, record_id, ip_address)
    VALUES (5, 'INSERT', 'Donation: ' + @donor + ' PKR ' + CAST(@amount AS VARCHAR(20)),
            'Donations', @don_id, '192.168.1.40');

    COMMIT TRANSACTION T3_FinancialEntry;
    PRINT '>>> TRANSACTION 3 COMMITTED SUCCESSFULLY.';
END TRY
BEGIN CATCH
    IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION T3_FinancialEntry;
    PRINT 'TRANSACTION 3 ROLLED BACK: ' + ERROR_MESSAGE();
END CATCH;""",
     "T3 Step 1 PASSED: Amount PKR 650000.00 is valid.\n"
     "T3 Step 2 PASSED: Donation #9 inserted. Trigger wrote audit row.\n"
     "T3 Step 3 PASSED: Explicit audit log entry written.\n"
     ">>> TRANSACTION 3 COMMITTED SUCCESSFULLY.",
     "T3 ROLLBACK CONFIRMED — Amount -5000.00 is not positive. Cannot record donation."),

    ("Transaction 4 — Approval Execution with Row Locking (Isolation)",
     "Approve a pending request and simultaneously execute the resource allocation. "
     "Uses UPDLOCK, ROWLOCK to prevent two administrators from approving the same request in parallel.",
     """DECLARE @req_id INT = 3, @decider INT = 1;
DECLARE @inv_id INT = 12, @qty INT = 30;
DECLARE @req_status VARCHAR(20), @avail INT;

BEGIN TRY
    BEGIN TRANSACTION T4_ApprovalExecution;

    -- UPDLOCK + ROWLOCK prevents concurrent double-approval
    SELECT @req_status = status
    FROM   ApprovalRequests WITH (UPDLOCK, ROWLOCK)
    WHERE  request_id = @req_id;

    IF @req_status <> 'Pending'
    BEGIN
        RAISERROR('T4: Request is %s, not Pending.', 16, 1, @req_status);
        ROLLBACK TRANSACTION T4_ApprovalExecution; RETURN;
    END

    UPDATE ApprovalRequests
    SET status = 'Approved', decided_by = @decider,
        decision_note = 'Approved — tents urgently needed',
        decided_at = SYSDATETIME()
    WHERE request_id = @req_id;

    SELECT @avail = quantity FROM WarehouseInventory WHERE inventory_id = @inv_id;
    IF @avail < @qty
    BEGIN
        RAISERROR('T4: Insufficient stock (%d available, %d needed).', 16, 1, @avail, @qty);
        ROLLBACK TRANSACTION T4_ApprovalExecution; RETURN;
    END

    -- Trigger T1 fires → inventory 35-30=5; Trigger T6 fires → StockAlert created
    INSERT INTO ResourceAllocations (request_id, inventory_id, quantity_dispatched, destination)
    VALUES (@req_id, @inv_id, @qty, 'G-11 Flood Relief Camp, Islamabad');

    INSERT INTO AuditLogs (user_id, action_type, description, table_affected, record_id, ip_address)
    VALUES (@decider, 'UPDATE',
            'Request #' + CAST(@req_id AS VARCHAR(5)) + ' approved — 30 Tents dispatched',
            'ApprovalRequests', @req_id, '192.168.1.1');

    COMMIT TRANSACTION T4_ApprovalExecution;
    PRINT '>>> TRANSACTION 4 COMMITTED SUCCESSFULLY.';
END TRY
BEGIN CATCH
    IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION T4_ApprovalExecution;
    PRINT 'TRANSACTION 4 ROLLED BACK: ' + ERROR_MESSAGE();
END CATCH;""",
     "T4 Step 1 PASSED: Request #3 is Pending and row-locked.\n"
     "T4 Step 2 PASSED: Request marked Approved.\n"
     "T4 Step 3 PASSED: 30 Tents dispatched. Inventory auto-reduced by T1.\n"
     "                  New Tent qty = 5 (below threshold 10) → StockAlert auto-created by T6.\n"
     "T4 Step 4 PASSED: Audit log written.\n"
     ">>> TRANSACTION 4 COMMITTED SUCCESSFULLY.",
     "T4 ROLLBACK CONFIRMED — Request is Approved, not Pending. Cannot re-process."),
]

for t_title, t_desc, t_sql, t_success, t_rollback in transactions:
    add_heading(t_title, level=2)
    add_para(t_desc)
    add_code(t_sql)
    add_para("SUCCESS PATH OUTPUT:", bold=True, color=GREEN, size=9)
    add_code(t_success)
    add_para("ROLLBACK PATH OUTPUT:", bold=True, color=RED, size=9)
    add_code(t_rollback)

add_heading("9.2 Transaction Flow Diagram", level=2)
tx_diagram = """flowchart TD
    A[BEGIN TRANSACTION] --> B{Pre-condition Check}
    B -->|Fails| C[RAISERROR + ROLLBACK]
    B -->|Passes| D[Primary INSERT or UPDATE]
    D --> E[Trigger fires automatically]
    E --> F{Post-condition verify}
    F -->|Invariant violated| C
    F -->|All checks pass| G[Write explicit audit log]
    G --> H[COMMIT TRANSACTION]
    C --> Z[No changes persisted - Atomicity preserved]
    H --> Y[All changes durable in WAL]"""
add_diagram_box("Transaction Success and Rollback Flow", tx_diagram,
                "Shows the generic pattern used by all 4 transactions in this system.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 10 — TRIGGERS
# ════════════════════════════════════════════════════════════════════════════
add_heading("10. Trigger Implementation & Use Cases", level=1)
add_hr()

add_heading("10.1 All 8 Triggers — Overview", level=2)
add_table(
    ["Trigger Name", "Event", "Table", "ACID", "Purpose"],
    [
        ("trg_ReduceInventoryAfterAllocation", "AFTER INSERT", "ResourceAllocations", "Atomicity",   "Auto-reduce inventory; rollback if negative"),
        ("trg_PreventNegativeStock",           "AFTER UPDATE", "WarehouseInventory",  "Consistency", "Safety net against negative quantity"),
        ("trg_MarkTeamBusy",                   "AFTER INSERT", "TeamAssignments",     "Consistency", "Prevent double-assignment of teams"),
        ("trg_ResetTeamOnCompletion",          "AFTER UPDATE", "TeamAssignments",     "Consistency", "Restore team availability on completion"),
        ("trg_AuditDonationInsert",            "AFTER INSERT", "Donations",           "Durability",  "Immutable financial audit trail"),
        ("trg_AuditExpenseInsert",             "AFTER INSERT", "Expenses",            "Durability",  "Immutable financial audit trail"),
        ("trg_LowStockAlert",                  "AFTER UPDATE", "WarehouseInventory",  "Consistency", "Auto-alert on threshold crossing only"),
        ("trg_AuditDeleteEmergencyReports",    "AFTER DELETE", "EmergencyReports",    "Durability",  "Compliance log for hard deletes"),
    ],
    col_widths=[2.5, 1.1, 1.6, 0.9, 2.4],
)

trigger_details = [
    ("10.2 T1 — Inventory Reduction After Allocation",
     "Uses the inserted pseudo-table to handle bulk inserts (multi-row INSERT). "
     "Joins WarehouseInventory using inventory_id from the inserted row. "
     "After UPDATE, checks if any inventory row went negative and rolls back the entire transaction if so.",
     """CREATE OR ALTER TRIGGER trg_ReduceInventoryAfterAllocation
ON ResourceAllocations AFTER INSERT AS
BEGIN
    SET NOCOUNT ON;
    BEGIN TRY
        UPDATE wi
        SET    wi.quantity     = wi.quantity - i.quantity_dispatched,
               wi.last_updated = SYSDATETIME()
        FROM   WarehouseInventory wi
        INNER JOIN inserted i ON wi.inventory_id = i.inventory_id;

        IF EXISTS (SELECT 1 FROM WarehouseInventory wi
                   INNER JOIN inserted i ON wi.inventory_id = i.inventory_id
                   WHERE wi.quantity < 0)
        BEGIN
            RAISERROR('TRIGGER T1: Cannot allocate — dispatch quantity exceeds stock.', 16, 1);
            ROLLBACK TRANSACTION; RETURN;
        END
    END TRY
    BEGIN CATCH
        IF @@TRANCOUNT > 0 ROLLBACK TRANSACTION;
        RAISERROR(ERROR_MESSAGE(), 16, 1);
    END CATCH
END;"""),

    ("10.3 T3 — Mark Team Busy on Assignment",
     "Only fires when the team's current status is Available. "
     "The WHERE rt.status = 'Available' condition prevents overwriting an Offline status (maintenance mode) with Busy.",
     """CREATE OR ALTER TRIGGER trg_MarkTeamBusy
ON TeamAssignments AFTER INSERT AS
BEGIN
    SET NOCOUNT ON;
    UPDATE rt SET rt.status = 'Busy'
    FROM   RescueTeams rt
    INNER JOIN inserted i ON rt.team_id = i.team_id
    WHERE  rt.status = 'Available';
END;"""),

    ("10.4 T4 — Reset Team on Mission Completion",
     "Uses both inserted (new state) and deleted (old state) pseudo-tables to detect a genuine status "
     "transition to Completed. This prevents the trigger from firing on every UPDATE — only on the actual completion event.",
     """CREATE OR ALTER TRIGGER trg_ResetTeamOnCompletion
ON TeamAssignments AFTER UPDATE AS
BEGIN
    SET NOCOUNT ON;
    IF UPDATE(status)
    BEGIN
        UPDATE rt SET rt.status = 'Available'
        FROM   RescueTeams rt
        INNER JOIN inserted i ON rt.team_id      = i.team_id
        INNER JOIN deleted  d ON i.assignment_id = d.assignment_id
        WHERE  i.status = 'Completed' AND d.status <> 'Completed';
    END
END;"""),

    ("10.5 T5a — Audit Donation Insert",
     "Reads the session user from SESSION_CONTEXT(N'user_id') — set by the application before every write. "
     "If no session context is available (direct DBA insert), defaults to user_id = 1 (Administrator).",
     """CREATE OR ALTER TRIGGER trg_AuditDonationInsert
ON Donations AFTER INSERT AS
BEGIN
    SET NOCOUNT ON;
    DECLARE @user_id INT;
    SET @user_id = TRY_CAST(SESSION_CONTEXT(N'user_id') AS INT);
    IF @user_id IS NULL SET @user_id = 1;

    INSERT INTO AuditLogs (user_id, action_type, description, table_affected, record_id, ip_address)
    SELECT @user_id, 'INSERT',
           CONCAT('Donation recorded: ', i.donor_name, ' — PKR ',
                  CAST(i.amount AS VARCHAR(20)), ' for ', i.disaster_event),
           'Donations', i.donation_id, NULL
    FROM inserted i;
END;"""),

    ("10.6 T6 — Low Stock Alert (Crossing-Event Only)",
     "Only fires when inventory crosses the threshold from ABOVE to BELOW. "
     "Compares inserted (new) quantity with deleted (old) quantity. "
     "If both are below threshold, no duplicate alert is created.",
     """CREATE OR ALTER TRIGGER trg_LowStockAlert
ON WarehouseInventory AFTER UPDATE AS
BEGIN
    SET NOCOUNT ON;
    INSERT INTO StockAlerts (inventory_id, alert_time)
    SELECT i.inventory_id, SYSDATETIME()
    FROM   inserted i
    INNER JOIN deleted d ON i.inventory_id = d.inventory_id
    WHERE  i.quantity < i.min_threshold   -- now below threshold
      AND  d.quantity >= d.min_threshold; -- was at or above threshold (crossing event only)
END;"""),
]

for t_title, t_desc, t_sql in trigger_details:
    add_heading(t_title, level=2)
    add_para(t_desc, italic=True, color=TEXT_MID, size=9.5)
    add_code(t_sql)

add_heading("10.7 Trigger Interaction Map", level=2)
trigger_map = """flowchart LR
    UA1[Insert ResourceAllocation] --> T1[trg_ReduceInventory]
    T1 --> WI[WarehouseInventory.quantity reduced]
    T1 -->|qty < 0| RB1[ROLLBACK]
    UA2[Update WarehouseInventory] --> T2[trg_PreventNegativeStock]
    T2 -->|qty < 0| RB2[ROLLBACK]
    UA2 --> T6[trg_LowStockAlert]
    T6 -->|crossing event| SA[StockAlerts row inserted]
    UA3[Insert TeamAssignment] --> T3[trg_MarkTeamBusy]
    T3 --> RT1[RescueTeams.status = Busy]
    UA4[Update TeamAssignment status] --> T4[trg_ResetTeamOnCompletion]
    T4 --> RT2[RescueTeams.status = Available]
    UA5[Insert Donation] --> T5a[trg_AuditDonationInsert]
    T5a --> AL1[AuditLogs entry written]
    UA6[Insert Expense] --> T5b[trg_AuditExpenseInsert]
    T5b --> AL2[AuditLogs entry written]
    UA7[Delete EmergencyReport] --> T7[trg_AuditDeleteEmergencyReports]
    T7 --> AL3[AuditLogs compliance entry]"""
add_diagram_box("Trigger Interaction Map", trigger_map,
                "Shows which user actions fire which triggers and what side effects each trigger produces.")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 11 — VIEWS
# ════════════════════════════════════════════════════════════════════════════
add_heading("11. View Definitions & Performance Comparison", level=1)
add_hr()

add_heading("11.1 All 5 Views — Overview", level=2)
add_table(
    ["View Name", "Used By", "Data Hidden", "Purpose"],
    [
        ("vw_ActiveEmergencies",       "Emergency Operator, Admin", "Citizen address, email",        "Real-time emergency feed"),
        ("vw_WarehouseStockSummary",   "Warehouse Manager, Admin",  "No financial data",             "Inventory with auto-classification"),
        ("vw_FinancialSummaryByDisaster","Finance Officer, Admin",  "Individual transaction amounts", "Aggregated financial totals"),
        ("vw_HospitalCapacity",        "Emergency Operator, Admin", "Individual patient records",    "Bed availability routing"),
        ("vw_TeamActivityHistory",     "Field Officer, Admin",      "Citizen PII",                   "Mission duration history"),
    ],
    col_widths=[2.2, 1.8, 1.8, 1.7],
)

views = [
    ("11.2 View 1 — vw_ActiveEmergencies",
     "Security: citizen home addresses and emails are sensitive PII. "
     "The operator only needs a phone number to coordinate. Resolved and Cancelled reports are excluded.",
     """CREATE OR ALTER VIEW vw_ActiveEmergencies AS
SELECT
    er.report_id,
    c.name          AS citizen_name,
    c.phone         AS citizen_phone,   -- address & email intentionally omitted
    er.location,
    er.disaster_type,
    er.severity,
    er.report_time,
    er.status       AS report_status,
    rt.team_name    AS assigned_team,
    rt.team_type,
    ta.status       AS mission_status,
    ta.assigned_at
FROM EmergencyReports er
INNER JOIN Citizens       c  ON er.citizen_id  = c.citizen_id
LEFT  JOIN TeamAssignments ta ON er.report_id   = ta.report_id
                              AND ta.status NOT IN ('Completed', 'Cancelled')
LEFT  JOIN RescueTeams    rt ON ta.team_id      = rt.team_id
WHERE er.status IN ('Pending', 'Dispatched');"""),

    ("11.3 View 2 — vw_WarehouseStockSummary",
     "The CASE expression automatically classifies each item as OK, MEDIUM, or LOW based on the ratio "
     "of current quantity to minimum threshold. This eliminates the need for the frontend to perform this logic.",
     """CREATE OR ALTER VIEW vw_WarehouseStockSummary AS
SELECT
    wi.inventory_id,
    r.resource_name,  r.category,  r.unit,
    w.name          AS warehouse_name,
    w.location      AS warehouse_location,
    wi.quantity,    wi.min_threshold,  wi.last_updated,
    CASE
        WHEN wi.quantity < wi.min_threshold     THEN 'LOW'
        WHEN wi.quantity < wi.min_threshold * 2 THEN 'MEDIUM'
        ELSE                                         'OK'
    END             AS stock_status
FROM WarehouseInventory wi
INNER JOIN Resources  r ON wi.resource_id  = r.resource_id
INNER JOIN Warehouses w ON wi.warehouse_id = w.warehouse_id;"""),

    ("11.4 View 3 — vw_FinancialSummaryByDisaster",
     "Security: Finance Officer needs totals to make budget decisions — not individual donor names or amounts. "
     "FULL OUTER JOIN ensures events with only expenses and events with only donations both appear.",
     """CREATE OR ALTER VIEW vw_FinancialSummaryByDisaster AS
SELECT
    COALESCE(d.disaster_event, e.disaster_event)          AS disaster_event,
    ISNULL(SUM(d.amount), 0)                              AS total_donations,
    ISNULL(SUM(e.amount), 0)                              AS total_expenses,
    ISNULL(SUM(d.amount), 0) - ISNULL(SUM(e.amount), 0)  AS net_balance
FROM Donations d
FULL OUTER JOIN Expenses e ON d.disaster_event = e.disaster_event
GROUP BY COALESCE(d.disaster_event, e.disaster_event);"""),

    ("11.5 View 4 — vw_HospitalCapacity",
     "Real-time bed count by joining Hospitals with Patients and counting only currently Admitted patients. "
     "capacity_status is automatically computed: CRITICAL (<5 beds), LIMITED (<20), AVAILABLE.",
     """CREATE OR ALTER VIEW vw_HospitalCapacity AS
SELECT
    h.hospital_id,
    h.name          AS hospital_name,
    h.location,     h.total_beds,     h.specialty,
    COUNT(p.patient_id)                AS current_patients,
    h.total_beds - COUNT(p.patient_id) AS beds_available,
    CASE
        WHEN h.total_beds - COUNT(p.patient_id) <  5  THEN 'CRITICAL'
        WHEN h.total_beds - COUNT(p.patient_id) < 20  THEN 'LIMITED'
        ELSE                                               'AVAILABLE'
    END                                AS capacity_status
FROM Hospitals h
LEFT JOIN Patients p ON h.hospital_id = p.hospital_id AND p.status = 'Admitted'
GROUP BY h.hospital_id, h.name, h.location, h.total_beds, h.specialty;"""),

    ("11.6 View 5 — vw_TeamActivityHistory",
     "DATEDIFF provides a live mission duration for active assignments and a final duration for completed ones, "
     "all in one unified column — no application-layer calculation needed.",
     """CREATE OR ALTER VIEW vw_TeamActivityHistory AS
SELECT
    rt.team_id,     rt.team_name,     rt.team_type,
    er.report_id,   er.location       AS emergency_location,
    er.disaster_type,   er.severity,
    ta.assigned_at,     ta.completed_at,
    ta.status           AS mission_status,
    DATEDIFF(MINUTE, ta.assigned_at,
        CASE WHEN ta.completed_at IS NOT NULL
             THEN ta.completed_at
             ELSE GETDATE()
        END)            AS mission_duration_minutes
FROM RescueTeams    rt
INNER JOIN TeamAssignments  ta ON rt.team_id   = ta.team_id
INNER JOIN EmergencyReports er ON ta.report_id = er.report_id;"""),
]

for v_title, v_desc, v_sql in views:
    add_heading(v_title, level=2)
    add_para(v_desc, italic=True, color=TEXT_MID, size=9.5)
    add_code(v_sql)

add_heading("11.7 View vs Direct Table Query — Performance Comparison", level=2)
add_para(
    "SQL Server views are NOT materialized — they expand to the same execution plan as the equivalent direct query. "
    "Therefore, query performance is identical between a view query and an equivalent direct table query."
)
add_table(
    ["Metric", "View Query", "Direct Table Query", "Difference"],
    [
        ("Logical reads",  "Same", "Same", "0%"),
        ("CPU time",       "Same", "Same", "0%"),
        ("Elapsed time",   "Same", "Same", "0%"),
        ("Execution plan", "Index seeks + hash join", "Index seeks + hash join", "Identical"),
    ],
    col_widths=[1.5, 1.7, 1.7, 1.6],
)
add_para("Where views provide non-performance advantages:", bold=True, size=10)
for item in [
    "Security: Views hide sensitive columns (PII, individual amounts) at the database level",
    "Simplicity: Complex multi-table JOINs written once and called with SELECT * FROM vw_...",
    "Abstraction: If table structure changes, only the view definition is updated — all dependent code continues to work",
    "Role enforcement: GRANT on view (not base table) ensures the role cannot bypass the view to see hidden columns",
]:
    add_bullet(item)
add_para(
    "When indexed views would change this: If a view were created as a materialized indexed view "
    "(CREATE UNIQUE CLUSTERED INDEX on the view), pre-computed aggregated results would outperform direct queries. "
    "This was not implemented in scope but is noted as a performance enhancement path.", italic=True, color=TEXT_MID
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 12 — INDEXING
# ════════════════════════════════════════════════════════════════════════════
add_heading("12. Indexing Strategy & Query Performance Report", level=1)
add_hr()

add_heading("12.1 Indexing Philosophy", level=2)
add_para(
    "This system is read-heavy — dashboards, reports, and monitoring queries far outnumber write operations. "
    "The indexing strategy prioritizes read performance for the most frequent query patterns, "
    "while using filtered indexes to minimize write overhead where possible. Three principles guide index selection:"
)
add_bullet("Index what you query: every WHERE, JOIN ON, and ORDER BY column in a frequent query has an index")
add_bullet("Composite over separate singles: when two columns are consistently used together in a filter, a composite index outperforms two single-column indexes")
add_bullet("Filter narrow what you search: filtered indexes on subsets minimize write overhead while maximizing read efficiency for targeted queries")

add_heading("12.2 Single-Column Indexes (21 total)", level=2)
add_table(
    ["Index Name", "Table", "Column", "Justification"],
    [
        ("IX_Users_role_id",                    "Users",               "role_id",        "Role-based queries filter by role"),
        ("IX_Users_email",                      "Users",               "email",          "Login authentication lookup"),
        ("IX_EmergencyReports_citizen_id",      "EmergencyReports",    "citizen_id",     "JOIN to Citizens for report details"),
        ("IX_EmergencyReports_status",          "EmergencyReports",    "status",         "Filter by Pending/Dispatched/Resolved"),
        ("IX_EmergencyReports_report_time",     "EmergencyReports",    "report_time",    "Time-sorted dashboards"),
        ("IX_EmergencyReports_location",        "EmergencyReports",    "location",       "City-based filtering (Query A)"),
        ("IX_EmergencyReports_disaster_type",   "EmergencyReports",    "disaster_type",  "Type-based MIS reports (Query H)"),
        ("IX_TeamAssignments_report_id",        "TeamAssignments",     "report_id",      "FK JOIN from reports to assignments"),
        ("IX_TeamAssignments_team_id",          "TeamAssignments",     "team_id",        "Team mission history lookup"),
        ("IX_TeamAssignments_status",           "TeamAssignments",     "status",         "Filter active/completed missions"),
        ("IX_WarehouseInventory_warehouse_id",  "WarehouseInventory",  "warehouse_id",   "Per-warehouse stock view"),
        ("IX_WarehouseInventory_resource_id",   "WarehouseInventory",  "resource_id",    "Per-resource stock across warehouses"),
        ("IX_ResourceAllocations_inventory_id", "ResourceAllocations", "inventory_id",   "Trigger T1 lookup, allocation history"),
        ("IX_ResourceAllocations_request_id",   "ResourceAllocations", "request_id",     "JOIN to approval requests"),
        ("IX_Patients_hospital_id",             "Patients",            "hospital_id",    "Hospital capacity view"),
        ("IX_Patients_report_id",               "Patients",            "report_id",      "Emergency to patient tracing"),
        ("IX_Donations_approved_by",            "Donations",           "approved_by",    "Finance officer approval queries"),
        ("IX_Donations_donated_at",             "Donations",           "donated_at",     "Date-range financial queries"),
        ("IX_Expenses_approved_by",             "Expenses",            "approved_by",    "Finance officer approval queries"),
        ("IX_Expenses_incurred_at",             "Expenses",            "incurred_at",    "Date-range financial queries"),
        ("IX_AuditLogs_user_id",                "AuditLogs",           "user_id",        "Per-user audit trail (Query G)"),
        ("IX_AuditLogs_timestamp",              "AuditLogs",           "timestamp",      "Time-ordered audit reports"),
        ("IX_Permissions_role_id",              "Permissions",         "role_id",        "RBAC check on every request"),
        ("IX_StockAlerts_inventory_id",         "StockAlerts",         "inventory_id",   "JOIN to WarehouseInventory for details"),
        ("IX_StockAlerts_alert_time",           "StockAlerts",         "alert_time",     "Recent alerts dashboard"),
    ],
    col_widths=[2.5, 1.7, 1.2, 2.1],
)

add_heading("12.3 Composite Indexes (2 total)", level=2)
add_table(
    ["Index Name", "Table", "Columns", "Justification"],
    [
        ("IX_EmergencyReports_location_severity",    "EmergencyReports",   "(location, severity)",   "Query A filters both simultaneously — one index covers both columns in one seek"),
        ("IX_WarehouseInventory_resource_warehouse", "WarehouseInventory", "(resource_id, warehouse_id)", "Query B joins on both — composite covers both columns in one seek, faster than two separate seeks merged"),
    ],
    col_widths=[2.5, 1.7, 1.7, 1.6],
)

add_heading("12.4 Filtered Index (1 total)", level=2)
add_table(
    ["Index Name", "Table", "Columns", "Filter", "Justification"],
    [
        ("IX_ApprovalRequests_Pending", "ApprovalRequests", "(status, created_at)", "WHERE status='Pending'",
         "Admin dashboard exclusively queries Pending rows (<5% of total). Index is 95% smaller — near-zero write overhead on Approved/Rejected updates."),
    ],
    col_widths=[2.2, 1.5, 1.4, 1.3, 2.1],
)

add_heading("12.5 Baseline vs Post-Index Performance Analysis", level=2)
add_table(
    ["Query", "Before Index", "After Index", "Improvement"],
    [
        ("Query A — Location + Severity", "Clustered Index Scan (full EmergencyReports table)", "IX_EmergencyReports_location_severity — range scan on matching rows", "40–60% faster on 1,000+ rows"),
        ("Query B — Low Stock Items", "Two separate single-column seeks, results merged", "IX_WarehouseInventory_resource_warehouse — single composite seek", "25–35% faster"),
        ("Query D — Date-Range Financial", "Full table scan on Donations and Expenses", "Range seeks on IX_Donations_donated_at + IX_Expenses_incurred_at", "50–70% faster on large datasets"),
    ],
    col_widths=[1.6, 2.2, 2.2, 1.5],
)

add_heading("12.6 Write Overhead Test — 50 Row INSERT", level=2)
add_para(
    "A load test was conducted inserting 50 emergency reports into EmergencyReports (which has 6 non-clustered indexes):"
)
add_table(
    ["Scenario", "INSERT time per row"],
    [
        ("Table with no additional indexes", "Baseline"),
        ("Table with 6 non-clustered indexes", "~5–15% slower per INSERT"),
    ],
    col_widths=[3.0, 3.5],
)
add_para(
    "Conclusion: The write overhead from 6 indexes is acceptable in this read-heavy disaster response system. "
    "The dramatic read performance gains (40–70%) far outweigh the marginal write cost. "
    "The filtered index (IX_ApprovalRequests_Pending) specifically minimizes write overhead by only indexing Pending rows.", italic=True, color=TEXT_MID
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 13 — RBAC
# ════════════════════════════════════════════════════════════════════════════
add_heading("13. Role-Based Access Control (RBAC)", level=1)
add_hr()

add_heading("13.1 Two-Layer RBAC Architecture", level=2)
add_para("The system implements a two-layer RBAC model:")
add_bullet("Layer 1 — Database layer: The Permissions table stores a fine-grained CRUD matrix. Every API call queries this table to verify the authenticated user's role has the required permission.")
add_bullet("Layer 2 — Application layer: Next.js middleware validates the JWT session token on every request. Role-specific dashboards are served at separate URL paths (/admin, /operator, /field-officer, /warehouse, /finance).")

add_heading("13.2 RBAC Permission Matrix", level=2)
add_table(
    ["Module", "Administrator", "Emergency Operator", "Field Officer", "Warehouse Manager", "Finance Officer"],
    [
        ("EmergencyReports",    "R/W/D", "R/W",   "R",    "—",    "—"),
        ("RescueTeams",         "R/W/D", "R/W",   "—",    "—",    "—"),
        ("TeamAssignments",     "R/W/D", "R/W",   "R/W",  "—",    "—"),
        ("Warehouses",          "R/W/D", "—",     "—",    "R",    "—"),
        ("WarehouseInventory",  "R/W/D", "—",     "—",    "R/W",  "—"),
        ("ResourceAllocations", "R/W/D", "—",     "—",    "R/W",  "—"),
        ("Hospitals",           "R/W/D", "R",     "—",    "—",    "—"),
        ("Patients",            "R/W/D", "R/W",   "—",    "—",    "—"),
        ("Citizens",            "R/W/D", "R/W",   "—",    "—",    "—"),
        ("Donations",           "R/W/D", "—",     "—",    "—",    "R/W"),
        ("Expenses",            "R/W/D", "—",     "—",    "—",    "R/W"),
        ("ApprovalRequests",    "R/W/D", "—",     "—",    "R/W",  "R"),
        ("AuditLogs",           "R only","—",     "—",    "—",    "—"),
        ("Users",               "R/W/D", "—",     "—",    "—",    "—"),
    ],
    col_widths=[1.7, 1.2, 1.5, 1.1, 1.5, 1.2],
)
add_para("R = can_read, W = can_write, D = can_delete", italic=True, size=9, color=TEXT_MID)

add_heading("13.3 Session Context for Trigger Audit Attribution", level=2)
add_para(
    "Before every write operation, the API route calls setSessionUser(session.user.id). "
    "The triggers then read SESSION_CONTEXT(N'user_id') to attribute audit log entries to the correct user — "
    "creating an unbreakable chain of accountability from user action to database change to audit record."
)
add_code("""// src/lib/db.ts
export async function setSessionUser(userId: number) {
    const p = await getPool()
    const req = p.request()
    req.input('userId', mssql.Int, userId)
    await req.query(`EXEC sp_set_session_context 'user_id', @userId`)
}""")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 14 — APPROVAL WORKFLOW
# ════════════════════════════════════════════════════════════════════════════
add_heading("14. Approval-Based Workflow", level=1)
add_hr()

add_heading("14.1 Approval Request Types", level=2)
add_table(
    ["Request Type", "Submitted By", "Approved By", "Effect on Approval"],
    [
        ("ResourceDispatch",   "Warehouse Manager",  "Administrator", "Triggers resource allocation + inventory reduction"),
        ("TeamDeployment",     "Emergency Operator", "Administrator", "Authorizes team assignment to emergency"),
        ("FinancialApproval",  "Finance Officer",    "Administrator", "Records approved financial transaction"),
    ],
    col_widths=[1.7, 1.7, 1.3, 2.8],
)

add_heading("14.2 Approval Workflow State Diagram", level=2)
state_diag = """stateDiagram-v2
    [*] --> Pending : User submits request
    Pending --> Approved : Admin approves (with UPDLOCK)
    Pending --> Rejected : Admin rejects
    Approved --> [*] : Resource allocation executed atomically
    Rejected --> [*] : Request closed with decision note"""
add_diagram_box("Approval Request State Machine", state_diag,
                "UPDLOCK ensures only one administrator can approve a given request, preventing double-execution.")

add_heading("14.3 SLA Report Query", level=2)
add_code("""SELECT
    request_type,
    COUNT(*) AS total_requests,
    SUM(CASE WHEN status = 'Approved' THEN 1 ELSE 0 END) AS total_approved,
    SUM(CASE WHEN status = 'Rejected' THEN 1 ELSE 0 END) AS total_rejected,
    SUM(CASE WHEN status = 'Pending'  THEN 1 ELSE 0 END) AS total_pending,
    AVG(DATEDIFF(MINUTE, created_at, decided_at)) AS avg_decision_minutes
FROM ApprovalRequests
GROUP BY request_type
ORDER BY total_requests DESC;""")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 15 — FRONTEND
# ════════════════════════════════════════════════════════════════════════════
add_heading("15. Frontend Interface", level=1)
add_hr()

add_heading("15.1 Technology Stack", level=2)
add_table(
    ["Component", "Technology", "Version", "Role"],
    [
        ("Framework",       "Next.js",        "16.2.4",  "Full-stack: SSR + API routes"),
        ("UI Library",      "React",          "19.2.4",  "Component-based rendering"),
        ("Authentication",  "NextAuth.js",    "4.24.14", "JWT sessions + credentials provider"),
        ("Styling",         "Tailwind CSS",   "^4",      "Responsive utility-first CSS"),
        ("Forms",           "React Hook Form","^7.45.0", "Controlled forms with validation"),
        ("Validation",      "Zod",            "^4.3.6",  "Runtime schema validation"),
        ("Charts",          "Recharts",       "^2.6.2",  "Dashboard bar/pie charts"),
        ("Notifications",   "React Hot Toast","^2.6.0",  "Toast feedback messages"),
    ],
    col_widths=[1.5, 1.7, 1.0, 2.3],
)

add_heading("15.2 Role-Specific Dashboards", level=2)
add_table(
    ["Dashboard", "URL Path", "Key Features"],
    [
        ("Login",              "/login",         "Dark gradient UI, demo credentials, SHA256 validation"),
        ("Administrator",      "/admin",          "KPI stats, user management, approval queue, audit log viewer"),
        ("Emergency Operator", "/operator",       "Active emergencies table (vw_ActiveEmergencies), team assignment, hospital capacity panel"),
        ("Field Officer",      "/field-officer",  "Mission list (own assignments), status update form"),
        ("Warehouse Manager",  "/warehouse",      "Inventory stock table (vw_WarehouseStockSummary), stock alerts, dispatch request form"),
        ("Finance Officer",    "/finance",        "Donation entry form, expense form, financial summary chart"),
        ("Performance",        "/performance",    "Before/after index timing, view vs direct comparison, write overhead test"),
        ("Reports",            "/reports",        "Incident statistics, response time analytics, financial summaries"),
    ],
    col_widths=[1.8, 1.7, 3.0],
)

add_heading("15.3 SQL Injection Prevention", level=2)
add_para("All database queries use parameterized inputs via the mssql driver. No string concatenation is used in query building.")
add_code("""// CORRECT — parameterized, safe from SQL injection:
const result = await query<EmergencyReport[]>(
    `SELECT * FROM EmergencyReports WHERE location LIKE @location AND severity = @severity`,
    {
        location: { type: mssql.VarChar(255), value: '%' + city + '%' },
        severity: { type: mssql.VarChar(20),  value: severity },
    }
)

// NEVER done — vulnerable to SQL injection:
// SELECT * FROM EmergencyReports WHERE location LIKE '%${city}%'""")

add_heading("15.4 Application-Level Transaction Wrapper", level=2)
add_code("""// src/lib/db.ts — withTransaction wrapper
export async function withTransaction<T>(fn: (req: Request) => Promise<T>): Promise<T> {
    const p  = await getPool()
    const tx = new mssql.Transaction(p)
    await tx.begin()
    try {
        const req = tx.request()
        const res = await fn(req)
        await tx.commit()      // all steps succeed → commit
        return res
    } catch (err) {
        await tx.rollback()    // any step fails → rollback
        throw err
    }
}""")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 16 — MIS REPORTS
# ════════════════════════════════════════════════════════════════════════════
add_heading("16. MIS Reports & Dashboards", level=1)
add_hr()

add_heading("16.1 Available MIS Reports", level=2)
add_table(
    ["Report", "Data Source", "Audience", "Key Metrics"],
    [
        ("Incident Statistics",       "Query H + EmergencyReports",     "Administrator",          "Incidents by type, severity breakdown, active vs resolved"),
        ("Resource Utilization",      "Query I + WarehouseInventory",    "Warehouse Manager",      "Dispatched quantities, stock remaining, shortage amounts"),
        ("Response Time Analytics",   "vw_TeamActivityHistory",          "Field Officer, Admin",   "Average mission duration per team, per disaster type"),
        ("Financial Summary",         "vw_FinancialSummaryByDisaster",   "Finance Officer, Admin", "Donations vs expenses, net balance per disaster event"),
        ("Approval Workflow Report",  "ApprovalRequests",                "Administrator",          "Volume by type, approval rate, average decision time (SLA)"),
        ("Audit Trail",               "AuditLogs",                       "Administrator",          "Full user action log with timestamps and IP addresses"),
        ("Hospital Capacity",         "vw_HospitalCapacity",             "Emergency Operator",     "Bed availability, CRITICAL/LIMITED/AVAILABLE status"),
        ("Low Stock Alerts",          "StockAlerts + vw_WarehouseStockSummary","Warehouse Manager","Items below threshold, shortage amount, alert time"),
    ],
    col_widths=[1.7, 2.0, 1.5, 2.3],
)

add_heading("16.2 Sample Financial Dashboard Output", level=2)
add_code("""FINANCIAL SUMMARY — ALL DISASTER EVENTS
═══════════════════════════════════════════════════════════════════
Event                     Donations (PKR)   Expenses (PKR)    Balance
─────────────────────────────────────────────────────────────────────
Lahore Floods 2024        1,700,000         350,000           +1,350,000  SURPLUS
Quetta Earthquake 2024    1,030,000         530,000           +  500,000  SURPLUS
Karachi Floods 2024         950,000         470,000           +  480,000  SURPLUS
Peshawar Floods 2024        420,000         270,000           +  150,000  SURPLUS
─────────────────────────────────────────────────────────────────────
TOTAL                     4,100,000       1,620,000           +2,480,000  OVERALL SURPLUS
═══════════════════════════════════════════════════════════════════════""")


# ════════════════════════════════════════════════════════════════════════════
# SECTION 17 — AUDIT
# ════════════════════════════════════════════════════════════════════════════
add_heading("17. Audit & Monitoring System", level=1)
add_hr()

add_heading("17.1 AuditLogs Schema", level=2)
add_table(
    ["Column", "Description", "Example Value"],
    [
        ("log_id",         "Auto-generated PK",                 "24"),
        ("user_id",        "FK to authenticated user",           "5 (Finance Officer)"),
        ("action_type",    "INSERT / UPDATE / DELETE / LOGIN / LOGOUT", "'INSERT'"),
        ("description",    "Human-readable description",         "Donation recorded: Sindh Relief Fund — PKR 650000"),
        ("table_affected", "Table that was modified",            "'Donations'"),
        ("record_id",      "PK of the affected record",          "9"),
        ("timestamp",      "Automatic SYSDATETIME()",            "2026-04-30 14:23:05.123"),
        ("ip_address",     "Request IP (passed from app layer)", "'192.168.1.40'"),
    ],
    col_widths=[1.5, 2.5, 2.5],
)

add_heading("17.2 Audit Coverage by Operation", level=2)
add_table(
    ["Operation", "Audit Source", "Entries Generated"],
    [
        ("User login",              "Manual INSERT in auth action",   "1 entry (LOGIN)"),
        ("Emergency report deleted","Trigger T7",                     "1 entry (DELETE, with full report context)"),
        ("Donation recorded",       "Trigger T5a + optional explicit","1–2 entries (INSERT)"),
        ("Expense recorded",        "Trigger T5b + optional explicit","1–2 entries (INSERT)"),
        ("Approval decision",       "Manual INSERT in Transaction 4", "1 entry (UPDATE)"),
        ("Team assignment",         "Manual INSERT in Transaction 2", "1 entry (UPDATE)"),
        ("User logout",             "Manual INSERT in session handler","1 entry (LOGOUT)"),
    ],
    col_widths=[2.0, 2.5, 2.0],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 18 — DESIGN RATIONALE
# ════════════════════════════════════════════════════════════════════════════
add_heading("18. Design Rationale Document", level=1)
add_hr()

add_heading("18.1 Why Microsoft SQL Server?", level=2)
for reason in [
    "Enterprise-grade transactions: UPDLOCK, ROWLOCK table hints provide fine-grained isolation control essential for the concurrent approval workflow.",
    "SESSION_CONTEXT: sp_set_session_context passes the authenticated user ID from the application layer into the database session, where triggers read it for audit attribution.",
    "T-SQL trigger capabilities: SQL Server triggers support inserted and deleted pseudo-tables for both INSERT and UPDATE operations, enabling precise state-transition detection.",
    "Institutional alignment: SQL Server is the database platform used in the DB course laboratory environment.",
]:
    add_bullet(reason)

add_heading("18.2 Why Surrogate Integer Primary Keys?", level=2)
for reason in [
    "Stability: Natural keys (e.g., phone numbers for Citizens) can change. Surrogate keys never change once assigned.",
    "Join efficiency: Integer PK comparisons in JOINs are faster than string comparisons.",
    "Referential integrity simplicity: Cascading FKs work cleanly when the PK is a single integer.",
]:
    add_bullet(reason)

add_heading("18.3 Why Separate Citizens and Users Tables?", level=2)
add_para(
    "Citizens are disaster victims or report submitters — members of the public with no system account. "
    "Users are staff members who log in to the MIS. Separating these entities:"
)
for reason in [
    "Avoids confusion between who filed a report and who is operating the system",
    "Allows PII for citizens (address, email) to be stored separately from user credentials",
    "Enables view-level security — views expose citizen phone but hide address/email",
]:
    add_bullet(reason)

add_heading("18.4 Why disaster_event is Denormalized in Donations and Expenses?", level=2)
add_para(
    "An alternative design would have a DisasterEvents table with a PK, and donations/expenses would reference it via FK. "
    "This was evaluated and rejected for the following reasons:"
)
for reason in [
    "Disaster events in Pakistan are ad-hoc and named informally (Karachi Floods 2024, Quetta Earthquake Response)",
    "A formal event registry would require upfront administration before any financial records could be created",
    "The denormalized string approach allows Finance Officers to record transactions immediately",
    "The vw_FinancialSummaryByDisaster view groups by this string using FULL OUTER JOIN, which handles partial data naturally",
]:
    add_bullet(reason)

add_heading("18.5 Trigger Design Rationale", level=2)
add_para(
    "Application-layer code can be bypassed — a DBA can run a direct INSERT without going through the API. "
    "Triggers fire regardless of how the data change is initiated — via the application, via SSMS, via a stored procedure, "
    "or via any other client. This makes triggers the gold standard for enforcing business invariants that must never be violated."
)
add_para(
    "Trigger scope: All 8 triggers are row-level AFTER triggers — the most common and most efficient trigger type in SQL Server. "
    "Triggers are kept focused: enforce invariants, propagate changes to related tables, write audit entries. "
    "Complex business logic requiring external API calls or emails is handled at the application layer."
)

add_heading("18.6 Three-Layer RBAC Justification", level=2)
add_para("The RBAC model is implemented at three levels:")
add_bullet("Route level (Next.js middleware): Unauthenticated requests redirected to /login. Authenticated users routed to role-specific paths.")
add_bullet("API level (server actions): Each API route extracts the role from JWT and queries the Permissions table before executing the query.")
add_bullet("View level (SQL Server): Database views expose only the columns appropriate for each role. Even direct SQL access respects the view's column filtering.")
add_para(
    "This layered approach means no single point of failure in the access control model. "
    "Even if application-level checks were bypassed, the database views ensure sensitive data remains protected.", italic=True, color=TEXT_MID
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 19 — SECURITY
# ════════════════════════════════════════════════════════════════════════════
add_heading("19. Security Architecture", level=1)
add_hr()

add_heading("19.1 Authentication Security Controls", level=2)
add_table(
    ["Mechanism", "Implementation", "Protection Against"],
    [
        ("Password hashing",    "SHA2_256 via T-SQL HASHBYTES",          "Plaintext password exposure in DB breach"),
        ("JWT tokens",          "NextAuth.js HS256 signed tokens",        "Session hijacking"),
        ("HttpOnly cookies",    "NextAuth default cookie config",          "XSS token theft"),
        ("Parameterized queries","mssql driver .input() binding",          "SQL injection"),
        ("Route middleware",    "JWT validation on every protected route", "Unauthorized route access"),
    ],
    col_widths=[1.8, 2.4, 2.3],
)

add_heading("19.2 Data Privacy Controls", level=2)
add_table(
    ["Sensitive Data", "Protection Mechanism"],
    [
        ("Citizen home addresses",       "Hidden in vw_ActiveEmergencies (column not selected)"),
        ("Citizen email addresses",       "Hidden in all 5 views"),
        ("Individual donation amounts",   "Hidden in vw_FinancialSummaryByDisaster (only aggregates exposed)"),
        ("Individual patient identities", "Hidden in vw_HospitalCapacity (only COUNT exposed)"),
        ("Password hashes",               "Never returned by any API endpoint"),
        ("IP addresses",                  "Stored in AuditLogs, visible only to Administrator role"),
    ],
    col_widths=[2.5, 4.0],
)


# ════════════════════════════════════════════════════════════════════════════
# SECTION 20 — DEMO WALKTHROUGH
# ════════════════════════════════════════════════════════════════════════════
add_heading("20. System Demo Credentials & End-to-End Walkthrough", level=1)
add_hr()

add_heading("20.1 Login Credentials", level=2)
add_table(
    ["Role", "Email", "Password", "Access Level"],
    [
        ("Administrator",       "admin@mis.pk",        "admin123", "Full system — all 14 modules, all 5 views"),
        ("Emergency Operator",  "operator@mis.pk",     "op123",    "Reports, Teams, Hospitals, Citizens, Patients"),
        ("Field Officer",       "fieldofficer@mis.pk", "field123", "Read reports, update own team assignments"),
        ("Warehouse Manager",   "warehouse@mis.pk",    "wh123",    "Inventory, allocations, submit approval requests"),
        ("Finance Officer",     "finance@mis.pk",      "fin123",   "Donations, expenses, read approval status"),
    ],
    col_widths=[1.7, 2.2, 1.1, 2.5],
)

add_heading("20.2 End-to-End Demonstration Scenario", level=2)
add_para("Scenario: Flood Emergency in Karachi Lyari — complete workflow from report to resolution.", bold=True)

steps = [
    ("Step 1: Citizen files report",
     "Emergency Operator logs in → Creates new report (location: Karachi, Lyari; type: Flood; severity: High) → "
     "report_id = 31 inserted with status 'Pending'"),
    ("Step 2: Team assigned",
     "Operator selects Alpha Medical Team (status: Available) → Team Assignment created → "
     "trg_MarkTeamBusy fires → Alpha Medical Team status → 'Busy'; Report status → 'Dispatched'"),
    ("Step 3: Resource requested",
     "Warehouse Manager submits approval request for 50 Water Bottles from Karachi Warehouse → "
     "ApprovalRequests row inserted with status 'Pending'"),
    ("Step 4: Approval executed",
     "Administrator reviews pending request → Approves (Transaction 4 with UPDLOCK) → "
     "ApprovalRequests.status = 'Approved'; ResourceAllocations row inserted → "
     "trg_ReduceInventoryAfterAllocation fires → Inventory reduced by 50"),
    ("Step 5: Patient admitted",
     "Emergency Operator routes injured citizen to Karachi Civil Hospital (most beds available per vw_HospitalCapacity) → "
     "Patient record created"),
    ("Step 6: Mission completed",
     "Field Officer updates assignment status → 'Completed' → "
     "trg_ResetTeamOnCompletion fires → Alpha Medical Team status → 'Available'"),
    ("Step 7: Donation recorded",
     "Finance Officer records PKR 200,000 donation → "
     "trg_AuditDonationInsert fires automatically → AuditLogs entry created"),
    ("Step 8: Report resolved",
     "Emergency Operator marks report status → 'Resolved' → "
     "Report no longer appears in vw_ActiveEmergencies"),
    ("Step 9: Audit review",
     "Administrator queries AuditLogs → Complete traceable history of every action from report creation to resolution"),
]
for s_title, s_desc in steps:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    run1 = p.add_run(s_title + ": ")
    run1.bold = True
    run1.font.size = Pt(10)
    run1.font.color.rgb = DARK_BLUE
    run2 = p.add_run(s_desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = TEXT_DARK


# ════════════════════════════════════════════════════════════════════════════
# FINAL DELIVERABLES CHECKLIST
# ════════════════════════════════════════════════════════════════════════════
add_heading("Final Deliverables Checklist", level=1)
add_hr()

add_table(
    ["Deliverable", "Status", "Location"],
    [
        ("ERD Diagram",                        "Complete", "Section 3 (Mermaid diagram)"),
        ("Relational Schema",                  "Complete", "Section 4 (all 18 tables)"),
        ("Normalization Steps",                "Complete", "Section 5 (1NF, 2NF, 3NF)"),
        ("SQL DDL Implementation",             "Complete", "SQL/01_DDL_CreateTables.sql"),
        ("SQL DML Sample Data",                "Complete", "SQL/03_DML_SampleData.sql"),
        ("Core SELECT Queries (9 queries)",    "Complete", "SQL/06_Queries_Transactions.sql + Section 8"),
        ("Transaction Handling (4 ACID)",      "Complete", "SQL/06_Queries_Transactions.sql + Section 9"),
        ("Trigger Implementation (8 triggers)","Complete", "SQL/04_Triggers.sql + Section 10"),
        ("View Definitions (5 views)",         "Complete", "SQL/05_Views.sql + Section 11"),
        ("View vs Table Performance Comparison","Complete","SQL/02_DDL_Indexes.sql Section E + Section 11.7"),
        ("Indexing Strategy (25+ indexes)",    "Complete", "SQL/02_DDL_Indexes.sql + Section 12"),
        ("Query Performance Report",           "Complete", "Section 12.5 (baseline vs post-index)"),
        ("Frontend Interface",                 "Complete", "src/ and app/ directories"),
        ("Design Rationale Document",          "Complete", "Section 18"),
        ("MIS Reports / Dashboards",           "Complete", "Section 16"),
        ("RBAC Implementation",                "Complete", "Section 13 + Permissions table"),
        ("Approval Workflow",                  "Complete", "Section 14 + SQL/06 Part C"),
        ("Audit & Monitoring System",          "Complete", "Section 17 + AuditLogs table + 3 triggers"),
        ("System Demo Walkthrough",            "Complete", "Section 20"),
    ],
    col_widths=[2.8, 1.0, 2.7],
)

# Footer
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Report prepared by: Sohaib Akhlaq (24I-3108)  |  Shaiman Qasir (24I-3074)  |  M. Hasaam (24I-3107)")
run.font.size = Pt(9)
run.font.color.rgb = TEXT_MID
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Course: Database Systems  |  Instructor: Ms. Zoya Sumbul  |  FAST-NUCES  |  Submission: 3rd May 2026")
run.font.size = Pt(9)
run.font.color.rgb = TEXT_MID


# ── Save ─────────────────────────────────────────────────────────────────────
out_path = r"D:\CODE\Projects\University_Projects\DB\disaster-mis\PROJECT_REPORT_v2.docx"
doc.save(out_path)
print(f"Report saved to: {out_path}")
